from copy import deepcopy

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml.ns import qn

from build_chuong4 import (
    add_block,
    add_figure_placeholder,
    add_heading,
    add_table,
    format_paragraph,
    set_run_font,
)


SOURCE = "CHUONG_4_VSL_viet_lai_cap_nhat.docx"
FINAL = "CHUONG_4_VSL_viet_lai_rut_gon.docx"


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(13)


def make_replacement_sections():
    tmp = Document()
    configure_document(tmp)

    add_figure_placeholder(tmp, 5, "Kết quả đo tốc độ phương tiện trên giao diện", "ảnh có tốc độ xe hiển thị cạnh phương tiện, ví dụ 68 km/h, 74 km/h hoặc 82 km/h.")
    add_heading(tmp, "4.3. Kết quả thực nghiệm và kịch bản kiểm thử")
    add_block(tmp, """
Mục đích của phần kiểm thử là xác định hệ thống có hoạt động đúng theo luồng sử dụng thực tế hay không, đồng thời kiểm tra các chức năng có liên kết đúng với nhau hay không. Việc kiểm thử không dừng ở việc xem chương trình có mở được hay không, mà cần theo dõi quá trình dữ liệu đi từ video hoặc camera đến kết quả nhận diện, đếm xe, đo tốc độ và đề xuất VSL.

Quy trình kiểm thử được thực hiện bằng cách khởi động chương trình, đăng nhập bằng tài khoản hợp lệ, chọn video hoặc camera, kiểm tra vùng ROI và vị trí hai vạch A/B, sau đó chạy nhận diện phương tiện. Người vận hành quan sát số xe trong vùng giám sát, kiểm tra tốc độ khi xe đi qua hai vạch, xem tốc độ VSL được đề xuất và lưu kết quả sau phiên chạy. Với các kịch bản có tham số điều khiển trên giao diện, người vận hành thay đổi mật độ, thời tiết hoặc chế độ thủ công để kiểm tra phản ứng của hệ thống.

Các kịch bản được chọn theo những chức năng có thể quan sát trực tiếp trên giao diện và có thể chụp ảnh màn hình minh họa. Những trường hợp chỉ khác nhau ở mức tham số mô phỏng, chẳng hạn tách riêng nhiều mức mưa, nhiều mức sương mù hoặc nhiều mức tỷ lệ xe nặng, không được tách thành các kịch bản độc lập nếu không có dữ liệu thực tế đi kèm. Cách rút gọn này giúp kết quả kiểm thử tập trung vào chức năng chính và tránh tạo cảm giác hệ thống đã có nhiều số liệu hiện trường hơn thực tế.
""")

    test_rows = [
        [1, "Đăng nhập hệ thống", "Nhập tài khoản và mật khẩu hợp lệ", "Vào được giao diện chính", "Kiểm tra bước truy cập ban đầu"],
        [2, "Chọn video/camera", "Chọn nguồn dữ liệu đầu vào", "Video hiển thị trên giao diện", "Kiểm tra OpenCV đọc dữ liệu"],
        [3, "Nhận diện phương tiện", "Chạy YOLO trên video", "Xe được khoanh bounding box và có nhãn loại xe", "Kiểm tra mô hình nhận diện"],
        [4, "Kiểm tra ROI", "Quan sát xe trong và ngoài ROI", "Chỉ xe trong ROI được đưa vào thống kê", "Kiểm tra vùng giám sát"],
        [5, "Đo tốc độ A/B", "Xe đi qua hai vạch A và B", "Hiển thị tốc độ ước lượng của xe", "Kết quả mang tính mô phỏng"],
        [6, "Mật độ cao", "Video có nhiều xe trong ROI", "VSL giảm so với khi đường ít xe", "Kiểm tra phản ứng theo mật độ"],
        [7, "Thời tiết xấu", "Chọn mưa hoặc sương mù trên giao diện", "VSL giảm thêm so với trời quang", "Kiểm tra ảnh hưởng thời tiết"],
        [8, "Chế độ thủ công", "Người vận hành nhập tốc độ VSL", "Hệ thống ưu tiên tốc độ nhập tay", "Kiểm tra quyền can thiệp người vận hành"],
    ]
    add_table(tmp, ["STT", "Kịch bản kiểm thử", "Nội dung kiểm tra", "Kết quả mong đợi", "Ghi chú"], test_rows, [0.35, 1.35, 1.75, 1.55, 1.20], 8.5, "Bảng 4.1. Các kịch bản kiểm thử chính của hệ thống")

    add_block(tmp, """
Ở kịch bản đăng nhập, hệ thống kiểm soát được người dùng trước khi cho truy cập giao diện chính. Đây là bước cần thiết vì các thông tin giám sát và tốc độ VSL không nên được thay đổi tự do bởi người không có quyền vận hành. Khi đăng nhập đúng, cửa sổ chính được mở và người vận hành có thể tiếp tục chọn nguồn dữ liệu.

Ở kịch bản chọn video hoặc camera, OpenCV đọc được nguồn dữ liệu đầu vào và hiển thị các khung hình trên giao diện. Đây là cơ sở để các chức năng phía sau hoạt động. Nếu nguồn dữ liệu không mở được, các bước nhận diện, đếm xe và tính VSL cũng không thể thực hiện, vì vậy việc kiểm tra đầu vào được thực hiện trước khi đánh giá thuật toán.

Ở kịch bản nhận diện phương tiện, YOLO phát hiện các nhóm xe được cấu hình và vẽ bounding box, nhãn loại xe cùng độ tin cậy lên video. Kết quả này không chỉ để người vận hành quan sát mà còn là dữ liệu đầu vào cho lọc ROI, đếm xe, phân loại xe nặng và tính mật độ. Khi model nhận diện không ổn định, các kết quả thống kê phía sau cũng bị ảnh hưởng.

Ở kịch bản ROI, hệ thống kiểm tra vị trí đại diện của phương tiện so với vùng giám sát trên mặt đường. Chỉ các xe nằm trong ROI mới được đưa vào thống kê, nhờ đó giảm khả năng đếm nhầm xe ở lề đường, làn ngược chiều hoặc vùng không liên quan. Kết quả kiểm tra ROI có thể quan sát trực tiếp bằng cách so sánh các xe nằm trong và ngoài đa giác đã vẽ.

Ở kịch bản đo tốc độ A/B, hệ thống ghi nhận thời điểm xe đi qua vạch A và vạch B rồi tính tốc độ theo công thức v = s / t × 3,6. Chức năng này thể hiện được nguyên lý tính tốc độ từ video, nhưng kết quả chỉ là ước lượng vì khoảng cách thực tế giữa hai vạch chưa được khảo sát ngoài hiện trường và camera chưa được hiệu chỉnh phối cảnh đầy đủ. Nếu xe bị che khuất hoặc mất tracking, kết quả cũng có thể sai.

Ở kịch bản mật độ cao, khi số xe trong ROI tăng, tốc độ VSL được điều chỉnh giảm so với trường hợp đường ít xe. Đây là phản ứng phù hợp với nguyên lý VSL vì dòng xe đông thường làm khoảng cách an toàn giảm và nguy cơ phanh gấp tăng. Cửa sổ trượt giúp số xe trung bình ổn định hơn, tránh cho VSL thay đổi chỉ vì dao động nhất thời giữa các frame.

Ở kịch bản thời tiết xấu, người vận hành chọn mưa hoặc sương mù trên giao diện và hệ thống giảm thêm VSL so với điều kiện trời quang. Mức điều chỉnh này nhằm phản ánh rủi ro do mặt đường trơn hoặc tầm nhìn thấp. Do thời tiết hiện được chọn thủ công, kết quả cần được hiểu là kiểm thử chức năng điều chỉnh theo tham số chứ chưa phải kết quả tự động nhận diện thời tiết từ camera.

Ở kịch bản thủ công, người vận hành nhập tốc độ VSL và hệ thống ưu tiên giá trị nhập tay. Điều này cho thấy người vận hành vẫn giữ vai trò kiểm soát trong những tình huống thuật toán chưa phản ánh đầy đủ hoặc cần điều chỉnh tạm thời. Khi triển khai thực tế, thao tác này cần có xác nhận, phân quyền và lưu nhật ký để hạn chế sai sót.
""")

    add_figure_placeholder(tmp, 6, "Kết quả nhận diện phương tiện và vùng ROI trong video thử nghiệm", "ảnh màn hình có video giao thông, bounding box quanh xe, nhãn loại xe và vùng ROI.")
    add_figure_placeholder(tmp, 7, "Kết quả đo tốc độ phương tiện qua hai vạch A/B", "ảnh có hai vạch A/B và tốc độ xe hiển thị trên giao diện.")
    add_figure_placeholder(tmp, 8, "Kết quả hiển thị tốc độ VSL trên giao diện hệ thống", "ảnh giao diện có ô tốc độ VSL đề xuất, mật độ giao thông, thời tiết và trạng thái hệ thống.")

    add_heading(tmp, "4.4. Đánh giá hệ thống và kết quả đạt được")
    add_block(tmp, """
Sau quá trình kiểm thử, hệ thống đã mô phỏng được luồng xử lý cơ bản của một hệ thống giám sát giao thông thông minh có đề xuất tốc độ giới hạn linh hoạt VSL. Chương trình có sản phẩm chạy thử trên máy tính cá nhân, có giao diện đăng nhập, giao diện giám sát, nguồn dữ liệu video hoặc camera, mô hình nhận diện và các ô thông tin phục vụ theo dõi. Kết quả này cho thấy các thành phần chính đã được kết nối thành một quy trình có đầu vào, xử lý và đầu ra, dù hệ thống chưa đạt mức triển khai ngoài hiện trường.

Về đọc dữ liệu, OpenCV mở được video hoặc camera mô phỏng và đưa từng frame lên giao diện. Đây là bước nền tảng vì toàn bộ nhận diện và phân tích đều phụ thuộc vào việc nhận dữ liệu liên tục. Khi nguồn video được chọn đúng, người vận hành có thể chạy, dừng và theo dõi quá trình xử lý trên cùng một luồng hình ảnh.

Về nhận diện phương tiện, YOLO phát hiện được các nhóm car, motorbike, truck, bus và bicycle trong điều kiện video phù hợp. Bounding box, nhãn phương tiện và độ tin cậy được hiển thị để người vận hành kiểm tra trực tiếp. Quan trọng hơn, kết quả nhận diện được sử dụng tiếp cho đếm xe, lọc ROI, nhận biết xe nặng và tính mật độ, chứ không chỉ dừng ở việc vẽ khung trên video.

Về ROI và đếm xe, hệ thống lọc được các phương tiện nằm trong vùng giám sát và thống kê số xe thuộc khu vực đó. Số xe trung bình được dùng để phân loại trạng thái giao thông, giúp thuật toán VSL có một đầu vào ổn định hơn so với số xe của một frame đơn lẻ. Cách xử lý này giảm ảnh hưởng của các đối tượng ngoài mặt đường và hạn chế việc VSL thay đổi quá nhanh.

Về đo tốc độ A/B, hệ thống thể hiện được nguyên lý ghi nhận thời gian xe đi qua hai vạch và tính tốc độ theo khoảng cách cấu hình. Kết quả hiển thị trên giao diện giúp người vận hành kiểm tra được tốc độ ước lượng của phương tiện. Tuy nhiên, đây vẫn là kết quả mô phỏng, chưa thể thay thế thiết bị đo chuyên dụng vì khoảng cách thực tế và phối cảnh camera chưa được hiệu chỉnh đầy đủ.

Về tính VSL, hệ thống kết hợp mật độ, thời tiết, tỷ lệ xe nặng và sự cố để đề xuất tốc độ. Các kiểm thử rút gọn cho thấy tốc độ có xu hướng giảm khi mật độ tăng hoặc khi điều kiện thời tiết xấu hơn. Chế độ thủ công cho phép người vận hành nhập tốc độ và can thiệp khi cần, nên quyết định cuối cùng vẫn nằm trong sự kiểm soát của con người.

Về giao diện và lưu trữ, phần mềm hiển thị được video, ROI, bounding box, số xe, mật độ, tốc độ VSL, thời tiết và trạng thái hệ thống. Các kết quả sau phiên chạy được ghi nhận bằng CSV hoặc HTML Report, gồm số xe, mật độ, thời tiết, sự cố, tỷ lệ xe nặng, tốc độ VSL và thời gian phân tích. Đây là cơ sở để nhóm kiểm tra lại phiên chạy và trình bày kết quả trong báo cáo.
""")

    result_rows = [
        ["Đọc video/camera", "Hệ thống đọc được video hoặc camera mô phỏng", "Tạo dữ liệu đầu vào cho xử lý"],
        ["Nhận diện phương tiện", "Phát hiện được phương tiện bằng YOLO", "Cung cấp dữ liệu cho đếm xe và VSL"],
        ["Xác định ROI", "Lọc phương tiện theo vùng giám sát", "Giảm đếm nhầm ngoài khu vực phân tích"],
        ["Đếm xe", "Thống kê số xe trong ROI", "Làm cơ sở phân loại mật độ"],
        ["Đo tốc độ", "Ước lượng tốc độ qua hai vạch A/B", "Bổ sung thông tin tốc độ phương tiện"],
        ["Tính VSL", "Đề xuất tốc độ theo mật độ, thời tiết, xe nặng và sự cố", "Thể hiện trọng tâm đề tài"],
        ["Giao diện", "Hiển thị kết quả trực quan", "Hỗ trợ người vận hành quan sát"],
        ["Điều khiển thủ công", "Cho phép nhập tốc độ thủ công", "Giữ vai trò kiểm soát của con người"],
        ["Lưu báo cáo", "Ghi nhận dữ liệu sau phiên chạy", "Phục vụ đánh giá kết quả thử nghiệm"],
    ]
    add_table(tmp, ["Nội dung đánh giá", "Kết quả đạt được", "Ý nghĩa"], result_rows, [1.45, 3.15, 1.67], 9, "Bảng 4.2. Tổng hợp kết quả đạt được")

    add_heading(tmp, "4.5. Thảo luận kết quả")
    add_block(tmp, """
Kết quả kiểm thử cho thấy hệ thống đã thể hiện đúng nguyên lý cơ bản của VSL: tốc độ giới hạn thay đổi theo tình trạng giao thông thay vì giữ cố định trong mọi trường hợp. Khi mật độ phương tiện trong ROI tăng, khoảng cách giữa các xe có xu hướng giảm và nguy cơ phanh gấp tăng, vì vậy VSL được điều chỉnh theo hướng thấp hơn. Đây là xu hướng chính cần kiểm tra trong đề tài.

Thời tiết xấu làm VSL giảm thêm so với điều kiện trời quang. Mưa có thể làm mặt đường trơn và tăng quãng đường phanh, còn sương mù làm giảm tầm nhìn của người lái. Trong phiên bản hiện tại, thời tiết được chọn trên giao diện nên kết quả chứng minh khả năng xử lý tham số của thuật toán, chưa chứng minh khả năng tự động nhận diện thời tiết từ hình ảnh.

Tỷ lệ xe nặng và sự cố cũng được đưa vào quyết định VSL để mô hình gần với thực tế hơn. Xe tải và xe buýt có khối lượng lớn, quãng đường phanh dài hơn và có thể ảnh hưởng đến dòng xe phía sau. Khi có sự cố, tốc độ cần giảm để cảnh báo từ xa và tạo thêm thời gian phản ứng. Tuy nhiên, các mức giảm hiện tại vẫn là luật mô phỏng, chưa được hiệu chỉnh bằng dữ liệu của một tuyến đường cụ thể.

Chế độ thủ công là phần cần thiết vì AI chỉ hỗ trợ ra quyết định, không nên thay thế hoàn toàn người vận hành trong bài toán an toàn giao thông. Người vận hành có thể kiểm tra hình ảnh, lý do đề xuất và nhập tốc độ khi phát hiện tình huống đặc biệt. Nếu kết nối với biển báo thật, thao tác này cần đi kèm xác nhận, phân quyền và nhật ký để bảo đảm việc can thiệp có thể kiểm soát.

Nhìn chung, kết quả hiện tại có ý nghĩa ở mức kiểm thử chức năng và mô phỏng logic. Các xu hướng VSL giảm theo mật độ, thời tiết xấu, xe nặng hoặc sự cố là phù hợp với mục tiêu đề tài, nhưng chưa thể xem là số liệu hiện trường. Để triển khai thực tế cần bổ sung camera thật, cảm biến thật, dữ liệu giao thông đủ lớn, khoảng cách A/B được khảo sát và quy trình xác nhận trước khi gửi lệnh đến biển báo.
""")

    return tmp


def insert_placeholder_before_caption(doc, number, description):
    caption = next(p for p in doc.paragraphs if p.text.startswith(f"Hình 4.{number}."))
    fragment = Document()
    p = fragment.add_paragraph()
    format_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, before=8, after=2, line=1.0, first_line=0)
    r = p.add_run(f"[CHÈN HÌNH 4.{number} TẠI ĐÂY]")
    set_run_font(r, size=13, bold=True, color=(31, 78, 121))
    caption._p.addprevious(deepcopy(p._p))


doc = Document(SOURCE)
replace = make_replacement_sections()
configure_document(doc)

insert_placeholder_before_caption(doc, 1, "")
insert_placeholder_before_caption(doc, 2, "")
insert_placeholder_before_caption(doc, 3, "")
insert_placeholder_before_caption(doc, 4, "")

# Update the old conclusion/location captions so they remain consistent after removing the 20-case table.
for p in doc.paragraphs:
    if p.text.startswith("Bảng 4.4. Đề xuất nhóm vị trí đặt biển báo VSL"):
        for run in list(p.runs):
            p._p.remove(run._r)
        r = p.add_run("Bảng 4.3. Đề xuất nhóm vị trí đặt biển báo VSL")
        set_run_font(r, size=12, bold=True)
    elif p.text.startswith("Qua 20 kịch bản thử nghiệm"):
        for run in list(p.runs):
            p._p.remove(run._r)
        r = p.add_run("Qua các kịch bản kiểm thử chính, VSL giảm khi mật độ tăng, giảm khi thời tiết xấu, giảm khi tỷ lệ xe nặng cao và giảm mạnh khi có sự cố. Chế độ thủ công giúp người vận hành kiểm soát hệ thống, phù hợp với quan điểm AI chỉ hỗ trợ ra quyết định chứ không thay thế hoàn toàn con người. Kết quả này thể hiện đúng xu hướng của VSL, nhưng vẫn cần được hiểu là kết quả mô phỏng theo các tham số đã cài đặt.")
        set_run_font(r, size=13)

start = next(p for p in doc.paragraphs if p.text.startswith("4.3."))
end = next(p for p in doc.paragraphs if p.text.startswith("4.6."))
body = doc._element.body
removing = False
for child in list(body):
    if child == start._p:
        removing = True
    if removing and child != end._p and child.tag != qn("w:sectPr"):
        body.remove(child)
    if child == end._p:
        break

# Insert the replacement section content immediately before the retained 4.6 heading.
for child in list(replace._element.body):
    if child.tag != qn("w:sectPr"):
        end._p.addprevious(deepcopy(child))

# Restore the location-related placeholders in the correct flow before 4.9.
figure_tail = Document()
configure_document(figure_tail)
add_figure_placeholder(figure_tail, 9, "Đề xuất vị trí đặt biển báo VSL theo cụm camera trên tuyến", "sơ đồ tuyến hoặc hình minh họa camera CAM_01, CAM_02, CAM_03 và biển báo VSL đặt trước vùng giám sát.")
add_figure_placeholder(figure_tail, 10, "Minh họa giảm tốc theo chuỗi biển báo VSL", "sơ đồ các biển báo liên tiếp hiển thị 100 km/h, 80 km/h, 60 km/h, 40 km/h trước khu vực ùn tắc hoặc sự cố.")
old_49 = next(p for p in doc.paragraphs if p.text.startswith("4.9."))
for child in list(figure_tail._element.body):
    if child.tag != qn("w:sectPr"):
        old_49._p.addprevious(deepcopy(child))

# Normalize paragraph spacing for the inserted and retained content.
for paragraph in doc.paragraphs:
    for run in paragraph.runs:
        if run.font.size is None:
            set_run_font(run, size=13)
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0

doc.save(FINAL)
print(FINAL)
