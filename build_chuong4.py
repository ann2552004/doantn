from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUT = "CHUONG_4_VSL_viet_lai_base.docx"


def set_run_font(run, name="Times New Roman", size=13, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def format_paragraph(p, align=None, before=0, after=6, line=1.15, first_line=0.3):
    pf = p.paragraph_format
    if align is not None:
        p.alignment = align
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if first_line is not None:
        pf.first_line_indent = Inches(first_line)


def add_para(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=0.3, after=6, italic=False):
    p = doc.add_paragraph()
    format_paragraph(p, align=align, after=after, first_line=first_line)
    r = p.add_run(text)
    set_run_font(r, size=13, italic=italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    format_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT, before=8 if level == 1 else 6, after=6, line=1.0, first_line=0)
    r = p.add_run(text)
    set_run_font(r, size=14 if level == 1 else 13, bold=True)
    return p


def add_title(doc, text):
    p = doc.add_paragraph()
    format_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=14, line=1.0, first_line=0)
    r = p.add_run(text)
    set_run_font(r, size=16, bold=True)
    return p


def add_equation(doc, text):
    p = doc.add_paragraph()
    format_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, before=4, after=4, line=1.0, first_line=0)
    r = p.add_run(text)
    set_run_font(r, size=13, italic=True)
    return p


def add_figure_placeholder(doc, number, title, description):
    p = doc.add_paragraph()
    format_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, before=8, after=2, line=1.0, first_line=0)
    r = p.add_run(f"[CHÈN HÌNH 4.{number} TẠI ĐÂY]")
    set_run_font(r, size=13, bold=True, color=(31, 78, 121))
    p2 = doc.add_paragraph()
    format_paragraph(p2, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=2, line=1.0, first_line=0)
    r2 = p2.add_run(f"Hình 4.{number}. {title}")
    set_run_font(r2, size=13, bold=True)
    p3 = doc.add_paragraph()
    format_paragraph(p3, align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=8, line=1.0, first_line=0)
    r3 = p3.add_run(f"Nội dung ảnh nên chèn: {description}")
    set_run_font(r3, size=12, italic=True, color=(89, 89, 89))


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_inches):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(sum(widths) * 1440)))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width * 1440)))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            set_cell_width(cell, widths[i])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_text(cell, text, size=10, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    format_paragraph(p, align=align, before=0, after=0, line=1.0, first_line=0)
    r = p.add_run(str(text))
    set_run_font(r, size=size, bold=bold)
    return p


def add_table(doc, headers, rows, widths, font_size=9, caption=None):
    if caption:
        p = doc.add_paragraph()
        format_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, before=6, after=4, line=1.0, first_line=0)
        r = p.add_run(caption)
        set_run_font(r, size=12, bold=True)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, size=font_size, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(table.rows[0].cells[i], "D9EAF7")
    for row_data in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            align = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[i], value, size=font_size, align=align)
    set_table_width(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_block(doc, text):
    for part in text.strip().split("\n\n"):
        if part.strip():
            add_para(doc, part.strip())


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.27)
section.page_height = Inches(11.69)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
normal.font.size = Pt(13)

add_title(doc, "CHƯƠNG 4. CÀI ĐẶT, THỬ NGHIỆM, ĐÁNH GIÁ VÀ KẾT LUẬN GIẢI PHÁP")

add_heading(doc, "4.1. Công nghệ sử dụng và môi trường cài đặt")
add_block(doc, """
Chương này trình bày phần cài đặt, thử nghiệm và đánh giá hệ thống giám sát giao thông thông minh có đề xuất tốc độ giới hạn linh hoạt VSL. Sau khi các chương trước đã giới thiệu lý do chọn đề tài, mục tiêu, cơ sở lý thuyết và nguyên lý xử lý, nội dung ở đây tập trung vào sản phẩm phần mềm đã xây dựng, luồng hoạt động thực tế, các kịch bản mô phỏng và mức độ đáp ứng của giải pháp. Cách trình bày được tổ chức theo đúng quá trình người vận hành sử dụng hệ thống, từ đăng nhập, chọn dữ liệu đầu vào, theo dõi phương tiện cho đến xem kết quả VSL.

Hệ thống được xây dựng ở mức mô phỏng trên máy tính cá nhân. Nguồn dữ liệu đầu vào gồm các video giao thông đã chuẩn bị sẵn hoặc camera được cấu hình trong chương trình. Khi một nguồn được chọn, chương trình đọc từng khung hình, đưa khung hình qua mô hình nhận diện, lọc đối tượng theo vùng giám sát ROI, cập nhật số lượng phương tiện và tính các thông số phục vụ đề xuất tốc độ. Kết quả xử lý được hiển thị trực tiếp trên giao diện PyQt5 để người vận hành có thể quan sát cả hình ảnh lẫn các giá trị phân tích.

Python được sử dụng làm ngôn ngữ lập trình chính vì phù hợp với việc kết hợp xử lý ảnh, mô hình học máy, giao diện desktop và lưu trữ dữ liệu trong cùng một chương trình. OpenCV đảm nhiệm việc mở video hoặc camera, đọc frame, chuyển đổi và vẽ thông tin lên khung hình như bounding box, nhãn phương tiện, ROI và hai vạch A/B. Mô hình YOLO được dùng để phát hiện phương tiện trong từng frame. Kết quả của YOLO gồm loại phương tiện, tọa độ bounding box và độ tin cậy, sau đó được chuyển cho các bước đếm xe, phân loại mật độ và theo dõi phương tiện.

PyQt5 đảm nhiệm phần giao diện, bao gồm màn hình đăng nhập, khu vực hiển thị video, các ô thông tin về số xe, mật độ, thời tiết, sự cố, tỷ lệ xe nặng và tốc độ VSL. NumPy hỗ trợ các phép tính trên tọa độ, danh sách điểm của ROI, bộ nhớ đệm số xe và các thao tác xử lý dữ liệu số. SQLite được sử dụng để lưu thông tin tài khoản và một số dữ liệu cấu hình cần thiết trong phạm vi mô phỏng. Dữ liệu kết quả có thể xuất thành CSV để thuận tiện xử lý bằng bảng tính, đồng thời tạo HTML Report để xem lại theo dạng báo cáo có cấu trúc. MQTT được đưa vào như một kênh mô phỏng việc gửi tốc độ VSL từ phần mềm đến thiết bị bên ngoài, chẳng hạn biển báo điện tử hoặc trung tâm điều khiển.

Việc lựa chọn các công nghệ trên phù hợp với phạm vi đồ án và điều kiện triển khai trên máy tính cá nhân. Python và OpenCV giúp rút ngắn thời gian phát triển phần xử lý video, YOLO đáp ứng yêu cầu nhận diện tương đối nhanh, còn PyQt5 giúp tạo được giao diện trực quan để trình diễn sản phẩm. SQLite, CSV/HTML Report và MQTT chưa biến hệ thống thành một nền tảng giao thông thực tế, nhưng chúng thể hiện được các thành phần cần có trong một luồng xử lý hoàn chỉnh: có người dùng, có dữ liệu, có kết quả phân tích, có lưu vết và có hướng truyền lệnh điều khiển.

Trong quá trình cài đặt, nhóm thực hiện nhận thấy hiệu năng của hệ thống phụ thuộc khá rõ vào cấu hình máy, đặc biệt là CPU, GPU, bộ nhớ và tốc độ đọc dữ liệu. Cùng một model YOLO nhưng khi chạy bằng CPU, số frame xử lý trong một giây có thể thấp hơn so với khi có GPU hỗ trợ. Tốc độ còn phụ thuộc vào kích thước ảnh đưa vào model, độ phân giải video, phiên bản thư viện và cách cài đặt môi trường Python. Vì vậy, kết quả trong chương này được hiểu là kết quả của một mô hình mô phỏng trên máy cá nhân, chưa phải thông số cam kết cho một hệ thống vận hành liên tục ngoài hiện trường.
""")
add_figure_placeholder(doc, 1, "Giao diện đăng nhập hệ thống", "ảnh màn hình đăng nhập của phần mềm, có ô tài khoản, mật khẩu và nút đăng nhập.")
add_figure_placeholder(doc, 2, "Giao diện chính của hệ thống giám sát giao thông và đề xuất VSL", "ảnh dashboard chính, có khung video, thông số giao thông, tốc độ VSL, nút chạy/dừng, chế độ tự động/thủ công và nhật ký hệ thống.")

add_heading(doc, "4.2. Các chức năng đã cài đặt")
add_block(doc, """
Các chức năng của hệ thống được liên kết thành một luồng xử lý từ nguồn dữ liệu đầu vào đến kết quả đề xuất tốc độ. Luồng này không chỉ cho thấy phần mềm có thể đọc và hiển thị video, mà còn thể hiện cách dữ liệu hình ảnh được chuyển dần thành thông tin giao thông có ý nghĩa cho người vận hành.

Chức năng đăng nhập được đặt ở bước đầu để kiểm soát quyền truy cập vào phần mềm. Người sử dụng cần nhập tài khoản và mật khẩu trước khi vào giao diện giám sát. Thông tin tài khoản được lưu trong cơ sở dữ liệu SQLite, nhờ đó chương trình có thể kiểm tra dữ liệu đăng nhập thay vì để bất kỳ người dùng nào mở trực tiếp màn hình điều khiển. Trong bối cảnh VSL, việc kiểm soát người dùng có ý nghĩa vì tốc độ đề xuất và tốc độ nhập thủ công đều liên quan đến vận hành giao thông. Ở phạm vi đồ án, đăng nhập chưa phải là một hệ thống phân quyền nhiều cấp, nhưng đã tạo được bước kiểm soát cơ bản và giúp sản phẩm gần với một phần mềm vận hành thực tế hơn.

Sau khi đăng nhập thành công, người vận hành có thể chọn video giao thông từ máy tính hoặc chọn một camera đã được cấu hình. Khi chọn video, OpenCV mở tệp, đọc các thông tin như kích thước khung hình và tốc độ khung hình, sau đó gửi từng frame đến chuỗi xử lý. Với camera, chương trình sử dụng cấu hình nguồn tương ứng để nhận dữ liệu liên tục. Việc cho phép sử dụng cả video và camera giúp hệ thống có thể kiểm thử lặp lại bằng cùng một dữ liệu, đồng thời vẫn giữ được hướng mở rộng sang giám sát trực tiếp.

Chức năng nhận diện phương tiện được thực hiện bằng mô hình YOLO. Trên mỗi frame, mô hình tìm các đối tượng phù hợp với các lớp đã cấu hình, gồm car, motorbike, truck, bus và bicycle. Kết quả của mỗi đối tượng có một bounding box xác định vùng chứa phương tiện, một nhãn loại xe và một giá trị độ tin cậy. Bounding box được vẽ lên video để người vận hành nhìn thấy vị trí phương tiện; nhãn xe được dùng cho việc phân loại; còn độ tin cậy giúp đánh giá mức chắc chắn của mô hình. Các phương tiện có độ tin cậy thấp có thể được loại khỏi thống kê theo ngưỡng cấu hình nhằm hạn chế nhận diện nhiễu.

Trong hệ thống, bounding box không chỉ có tác dụng minh họa. Từ bounding box, chương trình xác định điểm tâm hoặc một điểm đại diện ở phần dưới của xe để kiểm tra vị trí phương tiện trên mặt đường. Điểm này được dùng khi xét phương tiện có nằm trong ROI hay không và khi xác định xe đã đi qua vạch A hoặc vạch B. Việc thống nhất một điểm kiểm tra giúp giảm tình trạng chỉ cần một phần nhỏ của bounding box chạm vào ROI là xe đã bị tính nhầm vào vùng giám sát.

ROI, viết đầy đủ là Region of Interest, được dùng để giới hạn khu vực phân tích trên video. Camera thường ghi lại cả mặt đường, lề đường, dải phân cách, cây cối, biển báo hoặc làn xe không thuộc phạm vi cần thống kê. Nếu lấy toàn bộ khung hình để đếm, số liệu có thể bị ảnh hưởng bởi các đối tượng ngoài tuyến đang theo dõi. Vì vậy, hệ thống vẽ một đa giác ROI trên mặt đường và chỉ đưa các phương tiện có điểm đại diện nằm trong đa giác này vào bước thống kê. ROI cũng giúp việc đo tốc độ tập trung vào đúng đoạn đường đã được xác định giữa vạch A và vạch B.

Sau khi lọc theo ROI, chương trình đếm số phương tiện xuất hiện trong vùng giám sát. Số lượng trong một frame có thể dao động do xe vừa đi vào, vừa đi ra hoặc do model nhận diện khác nhau giữa các frame liên tiếp. Để tránh việc tốc độ VSL thay đổi quá nhanh, hệ thống lưu các giá trị đếm gần nhất trong một cửa sổ trượt rồi tính số xe trung bình. Giá trị trung bình này được dùng làm đầu vào ổn định hơn cho việc phân loại trạng thái giao thông. Khi số xe trung bình thấp, hệ thống có thể xem là đường thoáng; khi số xe tăng, trạng thái chuyển sang giao thông ổn định, mật độ cao hoặc nguy cơ ùn tắc tùy theo ngưỡng đã cấu hình.

Chức năng đo tốc độ được xây dựng theo phương pháp hai vạch A/B. Trên khung hình, người vận hành xác định vạch A ở vị trí xe bắt đầu được tính thời gian và vạch B ở vị trí xe kết thúc đoạn đo. Khi cùng một phương tiện đi qua A rồi đi qua B, chương trình lấy thời gian tương ứng giữa hai lần cắt vạch. Tốc độ được tính theo công thức:
""")
add_equation(doc, "v = s / t × 3,6")
add_block(doc, """
Trong đó, v là tốc độ của phương tiện, đơn vị km/h; s là khoảng cách thực tế giữa vạch A và vạch B, đơn vị mét; t là thời gian xe đi từ A đến B, đơn vị giây. Hệ số 3,6 dùng để đổi từ mét trên giây sang kilômét trên giờ. Nếu khoảng cách giữa hai vạch là 50 m và xe đi hết đoạn này trong 2 giây thì tốc độ ước lượng là v = 50 / 2 × 3,6 = 90 km/h. Trong video, thời gian có thể được suy ra từ số frame giữa hai lần qua vạch và FPS của nguồn dữ liệu. Kết quả đo được hiển thị cạnh phương tiện hoặc trong khu vực thông tin trên giao diện để người vận hành theo dõi.

Thuật toán VSL sử dụng số xe trung bình và trạng thái giao thông làm cơ sở tính tốc độ ban đầu. Từ mức tốc độ cơ sở, chương trình tiếp tục điều chỉnh theo thời tiết, tỷ lệ xe nặng, sự cố và trạng thái cảnh báo. Trời quang, ít xe và không có sự cố sẽ giữ tốc độ ở mức cao hơn. Mưa, sương mù, tỷ lệ truck/bus cao hoặc sự cố xảy ra sẽ làm tốc độ giảm để phản ánh điều kiện an toàn kém hơn. Các mức giảm trong đồ án được xây dựng theo luật mô phỏng, trong đó kết quả cuối cùng được giới hạn trong khoảng tốc độ tối thiểu và tối đa đã cài đặt để tránh sinh ra giá trị bất hợp lý.

Tỷ lệ xe nặng được tính từ số truck và bus trong số phương tiện nằm trong ROI. Xe nặng thường có khối lượng lớn, cần quãng đường phanh dài hơn và có thể làm dòng xe phía sau thay đổi. Vì vậy, khi tỷ lệ xe nặng vượt các ngưỡng cấu hình, hệ thống trừ thêm một mức tốc độ. Tương tự, sự cố được chia thành mức nhẹ và nghiêm trọng. Sự cố nhẹ tạo ra mức giảm để cảnh báo sớm, còn sự cố nghiêm trọng đưa VSL về gần mức tối thiểu nhằm buộc dòng xe phía sau giảm tốc mạnh hơn trước vùng nguy hiểm.

Ngoài chế độ tự động, hệ thống có chế độ thủ công. Ở chế độ này, người vận hành nhập tốc độ VSL mong muốn trên giao diện và chương trình ưu tiên giá trị đã nhập thay vì tiếp tục lấy kết quả tự động. Chức năng này cần thiết trong mô phỏng vì người vận hành có thể muốn can thiệp khi quan sát thấy sự kiện chưa được thuật toán phản ánh đầy đủ. Nếu mở rộng sang biển báo thật, giá trị thủ công nên đi kèm bước xác nhận và ghi nhật ký để tránh thay đổi tốc độ ngoài ý muốn.

Kết quả xử lý được lưu dưới dạng CSV hoặc HTML Report. CSV phù hợp với việc thống kê số xe, mật độ, tỷ lệ xe nặng, thời tiết, sự cố và VSL theo thời điểm. HTML Report giúp trình bày các phiên chạy dưới dạng dễ xem hơn, có thể mở bằng trình duyệt và dùng làm cơ sở đối chiếu khi đánh giá. Nhờ chức năng lưu báo cáo, quá trình thử nghiệm không chỉ dựa vào quan sát trên màn hình mà còn có dữ liệu để kiểm tra lại sau mỗi kịch bản.
""")
add_figure_placeholder(doc, 3, "Kết quả nhận diện phương tiện bằng YOLO", "ảnh video có bounding box quanh xe, có nhãn car, truck, bus, motorbike hoặc bicycle.")
add_figure_placeholder(doc, 4, "Vùng giám sát ROI và vạch đo tốc độ A/B trên video", "ảnh video có vùng ROI được vẽ trên mặt đường, có hai vạch A và B dùng để đo tốc độ.")
add_figure_placeholder(doc, 5, "Kết quả đo tốc độ phương tiện trên giao diện", "ảnh có tốc độ xe hiển thị cạnh phương tiện, ví dụ 68 km/h, 74 km/h, 82 km/h.")

add_heading(doc, "4.3. Kịch bản thử nghiệm và kết quả đạt được")
add_block(doc, """
Quy trình thử nghiệm được thực hiện theo luồng sử dụng của hệ thống. Trước hết, chương trình được khởi động trên máy tính và người vận hành đăng nhập bằng tài khoản đã lưu trong SQLite. Sau khi vào giao diện chính, người vận hành chọn một video giao thông hoặc camera đã cấu hình, kiểm tra khung hình hiển thị và xác nhận ROI đã phủ đúng phần mặt đường cần giám sát. Hai vạch A và B cũng được kiểm tra để bảo đảm nằm trên hướng di chuyển của xe. Khi các vùng cấu hình phù hợp, chế độ nhận diện được chạy và người vận hành quan sát bounding box, nhãn xe, số xe trong ROI, mật độ, tỷ lệ xe nặng và tốc độ VSL.

Sau bước chạy cơ bản, các tham số thời tiết và sự cố được thay đổi theo từng kịch bản. Nhóm thực hiện theo dõi sự thay đổi của số xe trung bình, trạng thái giao thông và tốc độ đề xuất, đồng thời thử cho xe đi qua hai vạch A/B để kiểm tra phần đo tốc độ. Chế độ tự động và thủ công được chạy riêng để kiểm tra sự khác nhau giữa tốc độ do thuật toán tính và tốc độ do người vận hành nhập.

Mục đích của thử nghiệm không chỉ là kiểm tra chương trình có mở được video và hiển thị hình ảnh hay không. Quan trọng hơn, thử nghiệm phải cho thấy thuật toán VSL phản ứng theo đúng logic đã thiết kế. Khi đường thông thoáng, VSL cần duy trì ở mức cao; khi số xe tăng, VSL phải giảm theo mật độ; khi mưa hoặc sương mù xuất hiện, tốc độ phải giảm thêm; khi xe nặng chiếm tỷ lệ cao hoặc có sự cố, tốc độ cần được hạ thấp để tăng khoảng thời gian phản ứng cho người lái. Với chế độ thủ công, hệ thống phải thể hiện rõ rằng tốc độ nhập từ người vận hành được ưu tiên.
""")

main_cases = [
    [1, "Trời quang, ít xe", "VSL cao"],
    [2, "Trời quang, nhiều xe", "VSL giảm theo mật độ"],
    [3, "Trời mưa", "VSL giảm thêm"],
    [4, "Sương mù", "VSL giảm mạnh hơn mưa"],
    [5, "Có nhiều xe tải/xe buýt", "VSL giảm do tỷ lệ xe nặng"],
    [6, "Có sự cố nhẹ", "VSL giảm để tăng an toàn"],
    [7, "Có sự cố nghiêm trọng", "VSL giảm mạnh"],
    [8, "Xe đi qua hai vạch A/B", "Hiển thị tốc độ xe"],
    [9, "Chế độ thủ công", "VSL lấy theo giá trị người vận hành nhập"],
]
add_table(doc, ["STT", "Kịch bản thử nghiệm", "Kết quả mong đợi"], main_cases, [0.45, 2.55, 3.15], 10, "Bảng 4.1. Các kịch bản thử nghiệm chính")

add_block(doc, """
Từ các kịch bản chính, nhóm thực hiện xây dựng thêm 20 trường hợp mô phỏng với các mức số xe trung bình, thời tiết, sự cố và tỷ lệ xe nặng khác nhau. Các giá trị VSL trong bảng là kết quả theo bộ ngưỡng mô phỏng của hệ thống, dùng để kiểm tra xu hướng thay đổi và mối quan hệ giữa các yếu tố đầu vào. Bảng không được dùng để khẳng định đây là tốc độ pháp lý áp dụng cho một tuyến cao tốc cụ thể; mục đích chính là cho thấy thuật toán phản ứng có quy luật và có thể giải thích được.
""")

scenario_rows = [
    [1, "Trời quang, ít xe", 3, "Trời quang", "Không", "10%", "91 km/h", "Đường thoáng, tốc độ giữ cao"],
    [2, "Trời quang, xe trung bình", 8, "Trời quang", "Không", "12%", "76 km/h", "Tốc độ giảm theo số xe"],
    [3, "Trời quang, nhiều xe", 13, "Trời quang", "Không", "15%", "56 km/h", "Mật độ cao, cần giảm tốc"],
    [4, "Nguy cơ ùn tắc", 18, "Trời quang", "Không", "18%", "40 km/h", "Chạm mức tốc độ tối thiểu"],
    [5, "Mưa nhỏ", 8, "Mưa nhỏ", "Không", "12%", "71 km/h", "Giảm nhẹ do mưa"],
    [6, "Mưa vừa", 8, "Mưa vừa", "Không", "12%", "66 km/h", "Mặt đường trơn hơn"],
    [7, "Mưa to", 8, "Mưa to", "Không", "12%", "61 km/h", "Tầm nhìn và độ bám giảm"],
    [8, "Sương mù mỏng", 8, "Sương mù mỏng", "Không", "12%", "66 km/h", "Giảm do tầm nhìn hạn chế"],
    [9, "Sương mù vừa", 8, "Sương mù vừa", "Không", "12%", "56 km/h", "Tầm nhìn hạn chế rõ"],
    [10, "Sương mù dày", 8, "Sương mù dày", "Không", "12%", "46 km/h", "Gần mức tốc độ tối thiểu"],
    [11, "Ít xe nặng", 8, "Trời quang", "Không", "15%", "76 km/h", "Chưa giảm do xe nặng"],
    [12, "Xe nặng trung bình", 8, "Trời quang", "Không", "25%", "72 km/h", "Giảm nhẹ do xe nặng"],
    [13, "Xe nặng cao", 8, "Trời quang", "Không", "40%", "68 km/h", "Giảm mạnh hơn"],
    [14, "Không sự cố", 10, "Trời quang", "Không", "15%", "70 km/h", "Hoạt động theo mật độ"],
    [15, "Sự cố nhẹ", 10, "Trời quang", "Nhẹ", "15%", "55 km/h", "Giảm để cảnh báo"],
    [16, "Sự cố nghiêm trọng", 10, "Trời quang", "Nghiêm trọng", "15%", "40 km/h", "Giảm mạnh về mức tối thiểu"],
    [17, "Xe qua vạch A/B", 6, "Trời quang", "Không", "10%", "82 km/h", "Có thể đo tốc độ xe"],
    [18, "Xe bị che khuất", 6, "Trời quang", "Không", "10%", "82 km/h", "Tốc độ có thể sai nếu mất tracking"],
    [19, "Chế độ tự động", 12, "Mưa vừa", "Không", "20%", "50 km/h", "VSL theo thuật toán"],
    [20, "Chế độ thủ công", 12, "Mưa vừa", "Không", "20%", "80 km/h", "Ưu tiên tốc độ người vận hành nhập"],
]
add_table(doc, ["STT", "Kịch bản", "Số xe TB", "Thời tiết", "Sự cố", "Tỷ lệ xe nặng", "VSL đề xuất", "Nhận xét"], scenario_rows, [0.30, 1.25, 0.50, 0.80, 0.70, 0.70, 0.70, 1.25], 8, "Bảng 4.2. Kết quả mô phỏng 20 kịch bản")

add_block(doc, """
Kết quả mô phỏng cho thấy mối quan hệ rõ giữa số xe trung bình và tốc độ VSL. Ở kịch bản chỉ có 3 xe trong ROI, hệ thống đề xuất 91 km/h vì đường còn thông thoáng. Khi số xe tăng lên 8 và 13 xe, tốc độ lần lượt giảm còn 76 km/h và 56 km/h. Với 18 xe, hệ thống chạm mức 40 km/h, thể hiện trạng thái có nguy cơ ùn tắc và cần ưu tiên an toàn hơn khả năng duy trì tốc độ cao.

Khi giữ số xe ở mức tương đối giống nhau nhưng thay đổi thời tiết, VSL tiếp tục giảm. Mưa nhỏ làm tốc độ giảm nhẹ xuống 71 km/h, mưa vừa xuống 66 km/h và mưa to còn 61 km/h. Với sương mù, mức giảm phụ thuộc vào mức độ hạn chế tầm nhìn: sương mù mỏng cho kết quả 66 km/h, sương mù vừa còn 56 km/h và sương mù dày còn 46 km/h. Điều này phù hợp với mục tiêu mô phỏng vì sương mù làm người lái khó quan sát chướng ngại vật và xe phía trước hơn.

Ảnh hưởng của xe nặng được thể hiện qua các kịch bản 11 đến 13. Khi tỷ lệ xe nặng là 15%, hệ thống chưa trừ thêm tốc độ do xe nặng. Khi tỷ lệ tăng lên 25%, VSL giảm nhẹ còn 72 km/h; khi tỷ lệ tăng lên 40%, tốc độ còn 68 km/h. Cách điều chỉnh này xuất phát từ việc xe tải và xe buýt có khối lượng lớn, thường cần quãng đường phanh dài hơn và làm dòng xe thay đổi khi xuất hiện với tỷ lệ cao.

Sự cố tạo ra mức giảm rõ rệt hơn so với trường hợp chỉ có mật độ. Ở trạng thái không sự cố, với 10 xe và 15% xe nặng, VSL là 70 km/h. Khi có sự cố nhẹ, tốc độ giảm còn 55 km/h để cảnh báo cho dòng xe phía sau. Khi sự cố nghiêm trọng, VSL giảm về 40 km/h, là mức tối thiểu trong mô phỏng. Kết quả này cho thấy trạng thái sự cố được ưu tiên trong quyết định VSL vì mức độ rủi ro không chỉ phụ thuộc vào số lượng xe.

Kịch bản xe qua hai vạch A/B cho thấy hệ thống có thể hiển thị tốc độ ước lượng khi phương tiện được theo dõi liên tục. Ngược lại, khi xe bị che khuất, tốc độ vẫn có thể được hiển thị theo giá trị tính được nhưng độ tin cậy giảm nếu tracking bị mất. Đây là điểm cần ghi nhận khi đánh giá kết quả: chức năng đo A/B thể hiện đúng nguyên lý, nhưng chất lượng số đo vẫn phụ thuộc vào khả năng duy trì dấu vết của cùng một phương tiện.

Hai kịch bản cuối cho thấy sự khác nhau giữa điều khiển tự động và thủ công. Ở chế độ tự động, với mưa vừa, 12 xe và 20% xe nặng, hệ thống đề xuất 50 km/h theo luật đã cài đặt. Khi chuyển sang thủ công và người vận hành nhập 80 km/h, giao diện hiển thị 80 km/h theo giá trị ưu tiên từ người vận hành. Điều này giúp người vận hành vẫn kiểm soát được hệ thống trong trường hợp cần điều chỉnh tạm thời, nhưng khi triển khai thật phải có quy trình xác nhận để hạn chế thao tác sai.
""")
add_figure_placeholder(doc, 6, "Các kịch bản thử nghiệm tiêu biểu của hệ thống VSL", "ảnh hoặc sơ đồ gồm các kịch bản trời quang ít xe, mật độ cao, mưa to, sương mù dày, nhiều xe nặng và có sự cố giao thông.")
add_figure_placeholder(doc, 7, "Bảng hoặc biểu đồ kết quả đề xuất tốc độ VSL theo kịch bản thử nghiệm", "bảng hoặc biểu đồ thể hiện VSL thay đổi theo các kịch bản, ví dụ 91 km/h khi ít xe, 56 km/h khi mật độ cao, 46 km/h khi sương mù dày, 40 km/h khi sự cố nghiêm trọng.")

add_block(doc, """
Về kết quả đạt được, hệ thống đã đọc được video giao thông và có khả năng tiếp nhận camera được cấu hình trong chương trình. Khung hình được đưa vào chuỗi xử lý liên tục, có thể chạy, tạm dừng và dừng theo thao tác của người vận hành. Đây là nền tảng để các chức năng nhận diện và phân tích hoạt động trên cùng một luồng dữ liệu.

Hệ thống đã tích hợp YOLO để nhận diện các nhóm phương tiện car, motorbike, truck, bus và bicycle. Bounding box, nhãn xe và độ tin cậy được vẽ trên khung hình, giúp kiểm tra trực tiếp kết quả của mô hình. Những đối tượng phù hợp với điều kiện lọc được đưa vào danh sách phương tiện phục vụ các bước sau.

ROI đã được cài đặt để lọc phương tiện theo vùng mặt đường cần giám sát. Sau khi lọc, hệ thống đếm được số xe trong ROI và tính số xe trung bình bằng cửa sổ trượt. Giá trị trung bình được dùng để phân loại mật độ, nhờ đó VSL không bị nhảy quá nhanh chỉ vì một frame có số lượng nhận diện khác biệt.

Chức năng đo tốc độ qua hai vạch A/B đã được mô phỏng và có thể hiển thị tốc độ ước lượng khi xe được theo dõi từ A đến B. Kết quả đo được sử dụng để quan sát hành vi của phương tiện và đối chiếu với tốc độ VSL. Do khoảng cách và phối cảnh còn mô phỏng, chức năng này được đánh giá ở mức thể hiện nguyên lý xử lý video, chưa phải phép đo thay thế thiết bị chuyên dụng.

Thuật toán VSL đã kết hợp được mật độ giao thông, thời tiết, tỷ lệ xe nặng, sự cố và trạng thái giao thông. Kết quả tốc độ được hiển thị trên giao diện cùng với các thông tin đầu vào và có thể chuyển sang chế độ thủ công. Hệ thống cũng lưu được dữ liệu phân tích dưới dạng CSV và HTML Report, tạo ra sản phẩm đầu ra có thể kiểm tra lại sau phiên chạy.
""")

summary_rows = [
    ["Đọc video/camera", "Mở và xử lý được nguồn video hoặc camera đã cấu hình", "Tạo dữ liệu đầu vào liên tục cho hệ thống giám sát"],
    ["Nhận diện phương tiện", "YOLO nhận diện car, motorbike, truck, bus, bicycle và hiển thị bounding box", "Cung cấp thông tin loại xe, vị trí và độ tin cậy"],
    ["Lọc theo ROI", "Chỉ đưa phương tiện nằm trong vùng giám sát vào thống kê", "Giảm ảnh hưởng của đối tượng ngoài khu vực cần theo dõi"],
    ["Đếm xe và mật độ", "Đếm xe trong ROI, tính số xe trung bình bằng cửa sổ trượt và phân loại trạng thái", "Tạo đầu vào ổn định cho thuật toán VSL"],
    ["Đo tốc độ A/B", "Ước lượng tốc độ dựa trên thời gian xe đi giữa hai vạch", "Minh họa khả năng trích xuất tốc độ từ video"],
    ["Tính VSL", "Điều chỉnh tốc độ theo mật độ, thời tiết, xe nặng, sự cố và trạng thái giao thông", "Đề xuất tốc độ linh hoạt theo tình huống"],
    ["Giao diện vận hành", "Hiển thị video, ROI, phương tiện, thông số và lý do đề xuất", "Giúp người vận hành quan sát và kiểm soát hệ thống"],
    ["Chế độ thủ công", "Cho phép nhập tốc độ và ưu tiên giá trị người vận hành nhập", "Duy trì khả năng can thiệp khi mô phỏng hoặc vận hành"],
    ["Lưu báo cáo", "Xuất dữ liệu CSV và HTML Report", "Phục vụ kiểm tra, đối chiếu và trình bày kết quả"],
]
add_table(doc, ["Nội dung đánh giá", "Kết quả đạt được", "Ý nghĩa"], summary_rows, [1.45, 3.15, 1.67], 9, "Bảng 4.3. Tổng hợp kết quả đạt được")
add_figure_placeholder(doc, 8, "Kết quả hiển thị tốc độ VSL trên giao diện hệ thống", "ảnh giao diện có ô tốc độ VSL đề xuất, ví dụ 50 km/h, 60 km/h, 70 km/h hoặc 80 km/h.")

add_heading(doc, "4.4. Hạn chế của hệ thống")
add_block(doc, """
Mặc dù hệ thống đã thể hiện được đầy đủ luồng xử lý chính, kết quả hiện tại vẫn chịu ảnh hưởng bởi nhiều giới hạn của dữ liệu và điều kiện mô phỏng. Những hạn chế này cần được nêu rõ để tránh hiểu rằng chương trình đã đạt mức sẵn sàng triển khai ngoài hiện trường.

Hạn chế dễ nhận thấy là chất lượng video đầu vào. Video bị mờ, rung, thiếu sáng hoặc có độ phân giải thấp làm cho phương tiện xuất hiện với kích thước nhỏ, đường biên không rõ và khó phân biệt với nền đường. Khi trời mưa, sương mù hoặc ánh sáng thay đổi, độ tương phản của xe giảm, còn nước trên ống kính có thể tạo ra vùng nhiễu. Xe bị che khuất bởi xe khác, bởi dải phân cách hoặc bởi góc quay không phù hợp cũng làm cho bounding box bị thiếu hoặc bị mất trong một số frame.

Mô hình YOLO đang sử dụng chưa được huấn luyện riêng cho điều kiện giao thông cao tốc Việt Nam. Các lớp car, truck và bus có thể bị nhầm trong trường hợp xe tải nhỏ, xe khách có hình dạng giống xe tải hoặc phương tiện bị nhìn từ xa. Việc phân loại sai loại xe ảnh hưởng trực tiếp đến tỷ lệ xe nặng. Nếu một xe con bị nhận thành truck hoặc bus, thuật toán VSL có thể giảm tốc nhiều hơn; ngược lại, nếu xe nặng bị bỏ sót, mức điều chỉnh an toàn có thể chưa đủ.

ROI hiện vẫn cần được chỉnh thủ công theo từng video hoặc từng camera. Khi ROI quá rộng, các xe ở làn ngoài, làn ngược chiều hoặc khu vực không thuộc đoạn cần phân tích có thể bị đưa vào thống kê. Khi ROI quá hẹp, phương tiện ở mép làn có thể bị bỏ sót. Ngoài ra, nếu camera thay đổi góc nhìn hoặc bị dịch chuyển, ROI cũ không còn phù hợp và cần được cấu hình lại.

Đo tốc độ bằng hai vạch A/B mới dừng ở mức ước lượng. Công thức v = s / t × 3,6 chỉ cho kết quả phù hợp khi khoảng cách s là khoảng cách thật giữa hai vị trí trên mặt đường và thời gian t được xác định đúng cho cùng một phương tiện. Video phối cảnh làm cho khoảng cách trên ảnh không tỷ lệ đều với khoảng cách ngoài thực tế, vì vậy việc chỉ dùng tọa độ pixel mà không hiệu chỉnh sẽ gây sai số. Kết quả còn phụ thuộc vào FPS và độ trễ khi đọc frame.

Khả năng tracking cũng ảnh hưởng đến tốc độ. Nếu xe bị che khuất trong thời gian ngắn, hai xe đi gần nhau hoặc detector thay đổi bounding box quá nhiều, chương trình có thể mất dấu phương tiện. Khi đó, thời điểm xe đi qua A và B có thể bị gán cho hai đối tượng khác nhau, làm tốc độ tính được sai. Kịch bản xe bị che khuất trong bảng thử nghiệm đã phản ánh giới hạn này.

Thời tiết hiện chủ yếu được người vận hành lựa chọn trên giao diện. Cách làm này thuận tiện cho thử nghiệm có kiểm soát, nhưng chưa phản ánh khả năng tự động của một hệ thống giám sát thật. Chương trình chưa tự suy ra mưa, sương mù hoặc mức tầm nhìn trực tiếp từ ảnh, cũng chưa nhận dữ liệu từ cảm biến mưa, độ ẩm hay tầm nhìn.

Thuật toán VSL được xây dựng theo các luật và ngưỡng giảm tốc mang tính mô phỏng. Các mức 5 km/h, 10 km/h, 15 km/h hoặc mức giảm do xe nặng và sự cố được chọn để kiểm tra xu hướng, chưa được hiệu chỉnh bằng dữ liệu lưu lượng, tốc độ, tai nạn và điều kiện mặt đường của một tuyến cụ thể. Vì vậy, tốc độ VSL trong bảng phù hợp để đánh giá logic phần mềm nhưng chưa thể dùng làm quy định vận hành.

Cuối cùng, hệ thống chưa kết nối với camera thật, cảm biến thời tiết thật và biển báo LED/VMS thật. MQTT hiện được dùng để mô phỏng hướng truyền lệnh, chưa kiểm tra đầy đủ độ tin cậy, xác thực, mất kết nối hay cơ chế an toàn của một mạng thiết bị ngoài thực địa. Giao diện desktop cũng mới phù hợp cho một máy trình diễn, chưa đáp ứng yêu cầu quản lý đồng thời nhiều tuyến và nhiều camera.
""")

add_heading(doc, "4.5. Giải pháp cải thiện và hướng phát triển")
add_block(doc, """
Các hướng phát triển được đề xuất dựa trên chính những hạn chế đã phân tích. Mục tiêu không chỉ là làm hệ thống chạy nhanh hơn, mà còn nâng độ tin cậy của dữ liệu đầu vào, làm rõ trách nhiệm của từng thành phần và đưa kết quả VSL gần hơn với điều kiện vận hành thực tế.

Trước hết, cần cải thiện camera và dữ liệu đầu vào. Camera nên có góc nhìn đủ cao, hướng nhìn rõ mặt đường, bao quát được các làn cần giám sát và hạn chế vùng bị che khuất. Độ phân giải, tốc độ khung hình và khả năng hoạt động ban đêm cần được lựa chọn phù hợp với khoảng cách quan sát. Bộ dữ liệu nên được thu thập ở nhiều điều kiện như trời quang, mưa nhỏ, mưa to, sương mù, ban ngày, ban đêm, mật độ thấp và mật độ cao để mô hình được kiểm tra trong nhiều tình huống hơn.

Mô hình YOLO cần được huấn luyện hoặc tinh chỉnh bằng dữ liệu giao thông Việt Nam. Dữ liệu nên có ô tô con, xe máy, xe tải nhẹ, xe tải nặng, xe khách, xe buýt, xe container và xe đạp với nhiều góc quay, khoảng cách và mức che khuất. Việc gán nhãn đúng giữa xe con và xe nặng sẽ giúp tỷ lệ xe nặng đáng tin cậy hơn, từ đó làm cho quyết định giảm tốc do xe nặng có cơ sở rõ hơn.

ROI nên được lưu riêng theo từng camera. Mỗi camera có thể có cấu hình gồm đa giác ROI, vị trí vạch A/B, khoảng cách thực tế giữa hai vạch và hướng di chuyển. Khi người vận hành đã cấu hình một camera, lần chạy sau chương trình tự nạp lại thay vì vẽ lại từ đầu. Ở mức cao hơn, có thể phát triển chức năng gợi ý ROI tự động dựa trên mặt đường, vạch làn hoặc vùng có mật độ đối tượng xuất hiện ổn định.

Đối với đo tốc độ, khoảng cách giữa A và B cần được khảo sát ngoài thực tế tại từng vị trí camera. Có thể dùng mốc đo trên mặt đường, bản vẽ tuyến hoặc dữ liệu bản đồ để xác định s. Sau đó áp dụng hiệu chỉnh phối cảnh, chẳng hạn biến đổi phối cảnh từ mặt phẳng ảnh sang mặt phẳng đường, để tọa độ xe phản ánh khoảng cách gần đúng hơn. Phần tracking cũng nên được nâng cấp bằng DeepSORT, ByteTrack hoặc OC-SORT để duy trì ID phương tiện ổn định hơn khi xe đi qua các frame liên tiếp.

Thời tiết có thể được tự động nhận diện bằng một mô hình AI riêng hoặc kết hợp giữa phân tích ảnh và cảm biến. Hệ thống có thể nhận dạng dấu hiệu mưa, sương mù, độ sáng thấp và tầm nhìn giảm từ camera; đồng thời nhận dữ liệu lượng mưa, độ ẩm, gió hoặc tầm nhìn từ cảm biến/API thời tiết. Khi có nhiều nguồn dữ liệu, thuật toán có thể kiểm tra chéo để giảm phụ thuộc vào lựa chọn thủ công của người vận hành.

Thuật toán VSL cần được hiệu chỉnh bằng dữ liệu thực tế của tuyến. Dữ liệu lịch sử về mật độ, tốc độ trung bình, tỷ lệ xe nặng, thời tiết, sự cố và tai nạn có thể được dùng để xác định ngưỡng phù hợp hơn. Ngoài điều khiển theo trạng thái hiện tại, hệ thống có thể bổ sung mô hình dự báo mật độ và tốc độ trong vài phút tiếp theo để giảm tốc sớm trước khi ùn tắc hình thành. Dù sử dụng luật hay mô hình học máy, kết quả vẫn cần được giới hạn trong các mức tốc độ được cơ quan vận hành cho phép.

Về giao diện, hệ thống có thể phát triển thành dashboard web để theo dõi nhiều camera, nhiều đoạn tuyến và nhiều biển báo trên cùng một màn hình. Dashboard có thể hiển thị bản đồ, trạng thái kết nối, lịch sử VSL, nhật ký sự cố và biểu đồ mật độ theo thời gian. Kiến trúc này phù hợp hơn với trung tâm điều hành so với giao diện desktop chỉ chạy trên một máy.

Khi kết nối với biển báo LED hoặc VMS thật, cần giữ cơ chế xác nhận của người vận hành trước khi gửi lệnh. Quy trình nên gồm bước hệ thống đề xuất tốc độ, hiển thị lý do và dữ liệu đầu vào, người vận hành kiểm tra, sau đó xác nhận để gửi lệnh qua MQTT hoặc giao thức phù hợp. Cần bổ sung xác thực thiết bị, ghi nhật ký, phát hiện mất kết nối và cơ chế đưa biển báo về trạng thái an toàn khi có lỗi. Điều này giúp hạn chế việc một sai số nhận diện hoặc một giá trị bất thường làm thay đổi biển báo ngoài ý muốn.
""")

add_heading(doc, "4.6. Đề xuất vị trí đặt biển báo VSL")
add_block(doc, """
Biển báo VSL phải được đặt đủ sớm trước khu vực có nguy cơ để người lái nhận biết và giảm tốc dần. Nếu biển đặt quá gần vị trí sự cố hoặc ùn tắc, người lái có thể phải phanh gấp, làm tăng nguy cơ va chạm phía sau. Ngược lại, nếu biển đặt quá xa mà không có biển nhắc lại, người lái có thể không còn nhớ tốc độ đã được thông báo khi đi đến vùng nguy hiểm. Vì vậy, khoảng cách đặt biển cần gắn với tốc độ khai thác, điều kiện tầm nhìn, độ dốc, độ cong và khả năng quan sát biển của tuyến.

Trên tuyến cao tốc, các vị trí cần ưu tiên gồm đầu tuyến, trước nút giao, trước lối ra hoặc lối vào, trước trạm thu phí, trước đoạn cầu, trước đoạn cong và trước khu vực có tầm nhìn hạn chế. Các đoạn thường xuyên ùn tắc, khu vực hay có mưa lớn hoặc sương mù cũng cần được xem xét bố trí VSL. Ngoài ra, khi một cụm camera phát hiện sự cố, biển báo nên được bố trí ở phía trước vùng giám sát để truyền thông tin cho các xe đang tiến đến chứ không chờ đến khi xe đã đi vào vị trí sự cố.

Trong mô hình mô phỏng, có thể xem CAM_01 đặt tại khoảng KM10, CAM_02 tại khoảng KM15 và CAM_03 tại khoảng KM20. Mỗi camera quan sát một vùng giao thông, còn biển VSL được đặt trước vùng giám sát khoảng 500 đến 1000 m tùy điều kiện tuyến. Khoảng cách này tạo ra thời gian phản ứng ban đầu cho người lái. Khi vùng giám sát phát hiện mật độ tăng hoặc có sự cố, tốc độ được truyền qua hệ thống để cập nhật biển báo ở phía trước vùng đó.
""")

location_rows = [
    ["Đầu tuyến cao tốc", "Xe bắt đầu vào tuyến và cần biết tốc độ khai thác hiện tại", "Thông báo tốc độ phù hợp ngay từ đầu"],
    ["Trước nút giao", "Có nhập làn, tách làn và chuyển làn", "Giảm xung đột giữa các dòng xe"],
    ["Trước lối ra/vào cao tốc", "Lưu lượng thay đổi, xe thường chuyển làn", "Ổn định tốc độ trước vùng giao cắt"],
    ["Trước trạm thu phí", "Xe thường giảm tốc hoặc dừng theo làn", "Hạn chế phanh gấp và va chạm phía sau"],
    ["Trước đoạn cầu", "Không gian xử lý và vị trí dừng xe hạn chế", "Tăng an toàn khi mật độ hoặc thời tiết xấu"],
    ["Trước đoạn cong", "Tầm nhìn và khả năng quan sát phía trước giảm", "Giảm nguy cơ mất lái và vào cua quá nhanh"],
    ["Khu vực tầm nhìn hạn chế", "Người lái khó nhận biết xe dừng hoặc chướng ngại vật", "Tạo cảnh báo sớm bằng tốc độ thấp hơn"],
    ["Khu vực thường xuyên ùn tắc", "Dòng xe giảm tốc bất thường theo thời điểm", "Điều hòa tốc độ từ xa, giảm dồn xe"],
    ["Khu vực hay mưa lớn/sương mù", "Mặt đường trơn hoặc tầm nhìn giảm", "Điều chỉnh tốc độ theo điều kiện quan sát"],
    ["Trước vùng CAM_01/CAM_02/CAM_03 tại KM10/KM15/KM20", "Camera phát hiện mật độ, xe nặng hoặc sự cố phía trước", "Đặt biển trước vùng giám sát khoảng 500–1000 m để xe kịp giảm tốc"],
]
add_table(doc, ["Nhóm vị trí", "Lý do đặt biển VSL", "Mục tiêu"], location_rows, [1.60, 2.50, 2.10], 9, "Bảng 4.4. Đề xuất nhóm vị trí đặt biển báo VSL")

add_block(doc, """
Trong tình huống ùn tắc hoặc sự cố nghiêm trọng, không nên chỉ đặt một biển ở sát điểm nguy hiểm. Có thể bố trí nhiều biển liên tiếp để giảm tốc theo chuỗi, chẳng hạn từ 100 km/h xuống 80 km/h, tiếp tục xuống 60 km/h và sau đó xuống 40 km/h trước vùng cần kiểm soát. Cách giảm tốc theo từng mức giúp dòng xe có thời gian điều chỉnh, hạn chế phanh đột ngột và làm cho tốc độ giữa các xe trong cùng dòng ổn định hơn. Khi sự cố được xử lý và dữ liệu giao thông trở lại bình thường, các biển có thể được cập nhật tăng dần theo quy trình xác nhận của người vận hành.
""")
add_figure_placeholder(doc, 9, "Đề xuất vị trí đặt biển báo VSL theo cụm camera trên tuyến", "sơ đồ tuyến hoặc hình minh họa camera CAM_01, CAM_02, CAM_03 và biển báo VSL đặt trước vùng giám sát.")
add_figure_placeholder(doc, 10, "Minh họa giảm tốc theo chuỗi biển báo VSL", "sơ đồ các biển báo liên tiếp hiển thị 100 km/h, 80 km/h, 60 km/h, 40 km/h trước khu vực ùn tắc hoặc sự cố.")

add_heading(doc, "4.7. Kết luận chương")
add_block(doc, """
Chương 4 đã trình bày quá trình cài đặt, thử nghiệm, đánh giá và đề xuất vị trí triển khai biển báo VSL cho hệ thống giám sát giao thông thông minh. Hệ thống được xây dựng bằng Python, sử dụng OpenCV để đọc và xử lý video, YOLO để nhận diện phương tiện, PyQt5 để xây dựng giao diện, NumPy để hỗ trợ xử lý dữ liệu số, SQLite để lưu tài khoản và cấu hình, CSV/HTML Report để lưu kết quả, còn MQTT được dùng để mô phỏng hướng truyền tốc độ đến biển báo điện tử.

Các chức năng đã cài đặt gồm đọc video hoặc camera, nhận diện car, motorbike, truck, bus và bicycle, hiển thị bounding box, nhãn và độ tin cậy, lọc phương tiện theo ROI, đếm xe, tính số xe trung bình, phân loại mật độ, đo tốc độ ước lượng qua hai vạch A/B, tính VSL theo mật độ, thời tiết, tỷ lệ xe nặng và sự cố, hiển thị kết quả trên giao diện, hỗ trợ chế độ thủ công và lưu báo cáo. Đây là một chuỗi xử lý tương đối đầy đủ cho một mô hình giám sát giao thông ở mức mô phỏng.

Kết quả 20 kịch bản cho thấy VSL giảm khi số xe tăng, giảm thêm khi thời tiết xấu, giảm khi tỷ lệ xe nặng cao và giảm mạnh khi có sự cố. Chế độ thủ công giúp người vận hành vẫn kiểm soát được tốc độ trong các tình huống cần can thiệp. Tuy nhiên, kết quả còn phụ thuộc vào chất lượng video, model YOLO, ROI, tracking và các ngưỡng mô phỏng; tốc độ A/B chưa được hiệu chỉnh ngoài thực tế; thời tiết chưa được nhận diện tự động; hệ thống chưa kết nối với camera, cảm biến và biển báo thật.

Hướng phát triển phù hợp là thu thập dữ liệu giao thông Việt Nam, huấn luyện lại YOLO, lưu ROI theo camera, hiệu chỉnh phối cảnh và khoảng cách A/B, nâng cấp tracking bằng DeepSORT, ByteTrack hoặc OC-SORT, tự động nhận diện thời tiết, hiệu chỉnh luật VSL bằng dữ liệu tuyến, phát triển dashboard web và kết nối LED/VMS có xác nhận của người vận hành. Với phạm vi hiện tại, hệ thống đã đạt mục tiêu mô phỏng nguyên lý kết hợp giữa thị giác máy tính và điều hành tốc độ, đồng thời tạo được nền tảng để tiếp tục phát triển theo hướng giao thông thông minh.
""")

add_title(doc, "KẾT LUẬN CHUNG")
add_block(doc, """
Đề tài “Xây dựng hệ thống giám sát giao thông thông minh và đề xuất tốc độ giới hạn linh hoạt VSL dựa trên thị giác máy tính” được thực hiện với mục tiêu mô phỏng một hệ thống có khả năng quan sát giao thông từ video hoặc camera và hỗ trợ đề xuất tốc độ theo tình trạng thực tế. Ý tưởng chính của đề tài là không xem tốc độ giới hạn như một giá trị cố định trong mọi thời điểm, mà thay đổi tốc độ theo mật độ phương tiện, điều kiện thời tiết, tỷ lệ xe nặng, sự cố và trạng thái giao thông.

Trong phạm vi đồ án, nhóm thực hiện đã xây dựng được phần mềm có khả năng đọc video hoặc camera, nhận diện phương tiện bằng YOLO, xác định vùng ROI, đếm xe, tính số xe trung bình, phân loại mật độ, đo tốc độ ước lượng qua hai vạch A/B và tính VSL theo bộ luật mô phỏng. Kết quả được đưa lên giao diện PyQt5 cùng các thông tin về bounding box, nhãn xe, độ tin cậy, mật độ, thời tiết, sự cố và tốc độ đề xuất. Hệ thống có chế độ tự động, chế độ thủ công, lưu dữ liệu CSV/HTML Report và MQTT để mô phỏng hướng truyền lệnh đến biển báo.

Về đóng góp cho lĩnh vực trí tuệ nhân tạo, đề tài đã ứng dụng mô hình nhận diện đối tượng vào một bài toán giao thông cụ thể. YOLO giúp chuyển hình ảnh video thành dữ liệu có cấu trúc gồm loại xe, vị trí xe và độ tin cậy. Dữ liệu này tiếp tục được sử dụng để đếm phương tiện, tính tỷ lệ xe nặng và tạo đầu vào cho thuật toán VSL. Qua đó, đề tài thể hiện được vai trò của thị giác máy tính trong việc trích xuất thông tin từ môi trường giao thông mà người vận hành khó có thể thống kê thủ công liên tục.

Về đóng góp cho giao thông thông minh, hệ thống đã mô phỏng được quy trình thu thập dữ liệu, phân tích trạng thái và hỗ trợ điều hành tốc độ. Kết quả thử nghiệm cho thấy khi mật độ tăng, thời tiết xấu, xe nặng chiếm tỷ lệ cao hoặc có sự cố, tốc độ đề xuất giảm theo hướng an toàn hơn. Khi đường thông thoáng và điều kiện quan sát tốt, tốc độ có thể giữ ở mức cao hơn. Đây là nguyên lý cơ bản của VSL và phù hợp với mục tiêu nghiên cứu của đề tài.

Hệ thống vẫn còn các hạn chế như phụ thuộc vào chất lượng video, chưa huấn luyện YOLO riêng cho dữ liệu cao tốc Việt Nam, ROI còn phải chỉnh thủ công, đo tốc độ A/B mới là ước lượng, tracking có thể mất dấu, thời tiết chủ yếu do người vận hành chọn và các ngưỡng VSL chưa được hiệu chỉnh bằng dữ liệu hiện trường. Hệ thống cũng chưa kết nối trực tiếp với camera thật, cảm biến thật và biển báo LED/VMS thật. Vì vậy, kết quả của đồ án cần được hiểu là kết quả ở mức mô phỏng, chưa phải giải pháp có thể đưa vào vận hành ngay.

Trong thời gian tới, đề tài có thể phát triển bằng cách thu thập dữ liệu giao thông thực tế, huấn luyện lại mô hình, tự động nhận diện thời tiết, hiệu chỉnh phối cảnh và tracking, xây dựng dashboard web cho nhiều camera, kết nối cảm biến/API thời tiết và tích hợp biển báo thật theo cơ chế xác nhận của người vận hành. Có thể bổ sung thêm chức năng dự báo mật độ và tốc độ trong vài phút tiếp theo để hệ thống chuyển từ phản ứng với hiện trạng sang hỗ trợ cảnh báo sớm.

Nhìn chung, đồ án đã hoàn thành mục tiêu đặt ra ở mức mô phỏng. Sản phẩm thể hiện được sự kết hợp giữa trí tuệ nhân tạo, xử lý ảnh, lập trình giao diện, lưu trữ dữ liệu và các nguyên lý của giao thông thông minh. Dù chưa thay thế được một hệ thống điều hành ngoài thực tế, kết quả đạt được là nền tảng phù hợp để tiếp tục nghiên cứu và hoàn thiện giải pháp giám sát giao thông và đề xuất tốc độ giới hạn linh hoạt VSL.
""")

# Apply a consistent font to all existing runs and set table paragraph spacing.
for paragraph in doc.paragraphs:
    for run in paragraph.runs:
        if run.font.size is None:
            set_run_font(run, size=13)
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0

doc.save(OUT)
print(OUT)
