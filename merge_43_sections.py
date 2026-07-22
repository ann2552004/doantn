from copy import deepcopy

from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn

from build_chuong4 import add_block, add_figure_placeholder, add_heading, add_table, format_paragraph, set_run_font


SOURCE = "CHUONG_4_VSL_cap_nhat_thoi_tiet.docx"
FINAL = "CHUONG_4_VSL_gop_muc_4_3.docx"


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(13)


def replace_paragraph(paragraph, text):
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)
    r = paragraph.add_run(text)
    set_run_font(r, size=13)


def make_section_43():
    tmp = Document()
    configure_document(tmp)
    add_heading(tmp, "4.3. Kiểm thử hệ thống và kết quả thực nghiệm")
    add_block(tmp, """
Sau khi hoàn thiện các chức năng chính, nhóm em tiến hành kiểm thử để đánh giá chương trình có hoạt động đúng theo mục tiêu đề tài hay không. Việc kiểm thử không chỉ nhằm kiểm tra chương trình có chạy được hay không, mà còn xem các chức năng có liên kết đúng với nhau, dữ liệu đầu vào có được xử lý ổn định và tốc độ giới hạn linh hoạt VSL (Variable Speed Limit) có thay đổi hợp lý theo từng tình huống giao thông hay không.

Trong phạm vi đồ án, hệ thống được kiểm thử cục bộ trên máy tính cá nhân sử dụng hệ điều hành Windows. Dữ liệu đầu vào là video giao thông và camera mô phỏng. Python được dùng làm ngôn ngữ lập trình chính; OpenCV đọc video, xử lý khung hình và hiển thị thông tin; YOLO (You Only Look Once) nhận diện phương tiện; PyQt5 xây dựng giao diện; SQLite lưu tài khoản và lịch sử đăng nhập; JSON (JavaScript Object Notation) lưu cấu hình camera, vùng ROI và vạch đo; CSV và HTML Report lưu kết quả; MQTT (Message Queuing Telemetry Transport) mô phỏng việc truyền tốc độ VSL đến biển báo điện tử. Do chưa kết nối camera, cảm biến và biển báo thật, kết quả trong mục này được hiểu là kết quả kiểm thử mô phỏng, chưa phải kết quả vận hành ngoài hiện trường.

Quy trình kiểm thử được thực hiện theo luồng sử dụng của hệ thống. Chương trình được khởi động, người vận hành đăng nhập, chọn video hoặc camera, kiểm tra khung hình hiển thị, kiểm tra vùng ROI (Region of Interest) và vị trí hai vạch A/B. Sau đó, hệ thống chạy nhận diện phương tiện bằng YOLO để người vận hành quan sát bounding box, nhãn xe và số lượng phương tiện trong ROI. Các bước tiếp theo gồm kiểm tra phân loại mật độ, đo tốc độ khi xe đi qua hai vạch A/B, theo dõi tốc độ VSL khi thay đổi mật độ hoặc thời tiết, kiểm tra chế độ thủ công và lưu kết quả ra báo cáo. Cách kiểm thử này giúp đánh giá toàn bộ luồng xử lý thay vì chỉ kiểm tra riêng từng nút bấm.
""")

    rows = [
        ["KT01", "Khởi động chương trình", "Chạy file chương trình chính", "Giao diện chương trình được mở", "Đạt"],
        ["KT02", "Đăng nhập hệ thống", "Nhập tài khoản và mật khẩu hợp lệ", "Vào được giao diện chính", "Đạt"],
        ["KT03", "Kiểm tra đăng nhập sai", "Nhập sai tài khoản hoặc mật khẩu", "Hệ thống không cho truy cập", "Đạt"],
        ["KT04", "Chọn video/camera", "Chọn nguồn dữ liệu đầu vào", "Video hiển thị trên giao diện", "Đạt"],
        ["KT05", "Nhận diện phương tiện", "Chạy YOLO trên video", "Phương tiện được khoanh bounding box và có nhãn loại xe", "Đạt"],
        ["KT06", "Kiểm tra ROI", "Quan sát xe trong và ngoài vùng ROI", "Chỉ xe trong ROI được đưa vào thống kê", "Đạt"],
        ["KT07", "Đếm xe trong ROI", "Chạy video có phương tiện đi qua vùng giám sát", "Số xe thay đổi theo khung hình", "Đạt"],
        ["KT08", "Phân loại mật độ", "So sánh đoạn ít xe và đoạn nhiều xe", "Trạng thái mật độ thay đổi theo số xe", "Đạt"],
        ["KT09", "Đo tốc độ qua hai vạch A/B", "Cho xe đi qua vạch A và vạch B", "Hiển thị tốc độ ước lượng của xe", "Đạt ở mức mô phỏng"],
        ["KT10", "Kiểm tra VSL theo mật độ", "Chạy đoạn video có nhiều xe trong ROI", "VSL giảm khi mật độ tăng", "Đạt"],
        ["KT11", "Kiểm tra VSL theo thời tiết", "Chọn mưa hoặc sương mù trên giao diện", "VSL giảm thêm so với trời quang", "Đạt"],
        ["KT12", "Kiểm tra xe nặng", "Quan sát trường hợp có xe tải hoặc xe buýt", "VSL có thể giảm thêm theo tỷ lệ xe nặng", "Đạt ở mức mô phỏng"],
        ["KT13", "Chế độ thủ công", "Nhập tốc độ VSL thủ công", "Hệ thống ưu tiên tốc độ người vận hành nhập", "Đạt"],
        ["KT14", "Lưu kết quả", "Xuất dữ liệu CSV hoặc HTML Report", "File kết quả được tạo và có thể xem lại", "Đạt"],
        ["KT15", "Xử lý lỗi đầu vào", "Chọn sai đường dẫn video hoặc nguồn không hợp lệ", "Hệ thống báo lỗi hoặc không bị treo", "Đạt một phần"],
        ["KT16", "Hiệu năng xử lý", "Chạy video có nhiều phương tiện", "Hệ thống vẫn xử lý nhưng phụ thuộc cấu hình máy", "Đạt một phần"],
    ]
    add_table(tmp, ["Mã kiểm thử", "Nội dung kiểm thử", "Cách thực hiện", "Kết quả mong đợi", "Kết quả đánh giá"], rows, [0.65, 1.35, 1.70, 1.70, 0.80], 7.5, "Bảng 4.1. Bảng kết quả kiểm thử hệ thống")

    add_block(tmp, """
Kết quả kiểm thử cho thấy phần đăng nhập hoạt động đúng, giúp kiểm soát người dùng trước khi vào giao diện giám sát. Khi nhập đúng tài khoản và mật khẩu, người vận hành truy cập được giao diện chính; khi nhập sai, hệ thống không cho phép truy cập tiếp. Đây là bước kiểm soát ban đầu cần thiết vì người vận hành có thể theo dõi và can thiệp vào tốc độ VSL.

Ở bước chọn video hoặc camera, hệ thống đọc được nguồn dữ liệu đầu vào và hiển thị khung hình trên giao diện. Đây là nền tảng cho các chức năng phía sau. Khi YOLO được chạy, phương tiện được phát hiện và hiển thị bằng bounding box, nhãn loại xe cùng độ tin cậy. Kết quả nhận diện không chỉ có ý nghĩa minh họa mà còn được dùng cho các bước lọc ROI, đếm xe, phân loại mật độ và nhận biết xe nặng.

Vùng ROI giúp hệ thống chỉ đưa các phương tiện nằm trong khu vực giám sát vào thống kê, hạn chế đếm nhầm xe ở lề đường, làn ngược chiều hoặc vùng không liên quan. Khi số xe trong ROI thay đổi, trạng thái mật độ cũng thay đổi theo. Việc sử dụng số xe trung bình bằng cửa sổ trượt giúp kết quả ổn định hơn và hạn chế tình trạng VSL thay đổi liên tục theo từng frame.

Chức năng đo tốc độ qua hai vạch A/B thể hiện được nguyên lý tính tốc độ từ video. Khi xe đi qua vạch A và tiếp tục đi qua vạch B, hệ thống ghi nhận thời gian di chuyển và tính tốc độ theo công thức v = s / t × 3,6. Kết quả này chỉ được đánh giá ở mức mô phỏng vì khoảng cách thực tế giữa hai vạch chưa được khảo sát ngoài hiện trường, camera chưa được hiệu chỉnh phối cảnh đầy đủ và tracking có thể mất dấu khi xe bị che khuất.

Đối với VSL, kiểm thử cho thấy tốc độ đề xuất thay đổi theo xu hướng hợp lý: khi mật độ tăng hoặc thời tiết xấu như mưa và sương mù, tốc độ đề xuất giảm; khi có xe tải, xe buýt hoặc sự cố, hệ thống có thể điều chỉnh giảm thêm. Các trường hợp thời tiết chỉ được trình bày ở mức đại diện, không tách thành nhiều mức nhỏ nếu không có dữ liệu thực tế hoặc ảnh giao diện đi kèm. Chế độ thủ công vẫn cho phép người vận hành nhập tốc độ để can thiệp khi cần.

Chức năng lưu kết quả giúp ghi nhận dữ liệu sau phiên chạy dưới dạng CSV hoặc HTML Report. Nhờ đó, nhóm em có thể xem lại số xe, mật độ, thời tiết, sự cố, tỷ lệ xe nặng, tốc độ VSL và thời gian phân tích. Hai trường hợp đạt một phần liên quan đến xử lý lỗi đầu vào và hiệu năng được ghi nhận một cách thận trọng vì kết quả còn phụ thuộc vào đường dẫn nguồn, chất lượng video, CPU/GPU và độ phân giải dữ liệu.
""")

    add_figure_placeholder(tmp, 6, "Kết quả kiểm thử nhận diện phương tiện và vùng ROI", "ảnh giao diện có video giao thông, bounding box quanh xe, nhãn loại xe và vùng ROI.")
    add_figure_placeholder(tmp, 7, "Kết quả kiểm thử đo tốc độ qua hai vạch A/B", "ảnh có hai vạch A/B và tốc độ xe hiển thị trên giao diện.")
    add_figure_placeholder(tmp, 8, "Kết quả kiểm thử tốc độ VSL trên giao diện", "ảnh giao diện có ô tốc độ VSL đề xuất, số xe, mật độ, thời tiết và trạng thái hệ thống.")

    add_block(tmp, """
Qua quá trình kiểm thử, các chức năng chính của hệ thống đã hoạt động đúng theo mục tiêu đề tài ở mức mô phỏng. Hệ thống có thể đọc video hoặc camera, nhận diện phương tiện, lọc ROI, đếm xe, phân loại mật độ, đo tốc độ ước lượng và đề xuất tốc độ VSL. Một số nội dung như đo tốc độ, thời tiết, xe nặng, sự cố và hiệu năng vẫn cần được hiểu là kiểm thử mô phỏng, chưa phải kết quả triển khai ngoài hiện trường. Các trường hợp chỉ đạt một phần chủ yếu liên quan đến chất lượng video, hiệu năng máy tính, cấu hình ROI, xử lý nguồn không hợp lệ và độ ổn định của tracking.
""")
    return tmp


doc = Document(SOURCE)
configure_document(doc)
fragment = make_section_43()

start = next(p for p in doc.paragraphs if p.text.startswith("4.3."))
end = next(p for p in doc.paragraphs if p.text.startswith("4.4."))
body = doc._element.body
removing = False
for child in list(body):
    if child == start._p:
        removing = True
    if removing and child != end._p and child.tag != qn("w:sectPr"):
        body.remove(child)
    if child == end._p:
        break

for child in list(fragment._element.body):
    if child.tag != qn("w:sectPr"):
        end._p.addprevious(deepcopy(child))

# Remove old numeric examples from the discussion so it remains consistent with the merged test section.
for p in doc.paragraphs:
    if p.text.startswith("Trong các yếu tố đầu vào, mật độ phương tiện là nền tảng quan trọng nhất"):
        replace_paragraph(p, "Trong các yếu tố đầu vào, mật độ phương tiện là nền tảng quan trọng nhất của thuật toán. Số xe trung bình trong ROI phản ánh mức độ chiếm dụng mặt đường và khoảng cách giữa các xe. Khi mật độ tăng, khoảng cách an toàn có xu hướng giảm và nguy cơ phanh gấp tăng, vì vậy VSL được điều chỉnh theo hướng thấp hơn. Đây là xu hướng phù hợp với nguyên lý của hệ thống, nhưng các mức giảm cụ thể vẫn mang tính mô phỏng.")
    elif p.text.startswith("Nhìn chung, kết quả hiện tại có ý nghĩa ở mức mô phỏng"):
        replace_paragraph(p, "Nhìn chung, kết quả hiện tại có ý nghĩa ở mức mô phỏng và kiểm tra logic. Tốc độ VSL thay đổi theo mật độ, thời tiết, xe nặng và sự cố, nhưng chưa thể xem là tốc độ pháp lý áp dụng cho một tuyến cụ thể. Để đánh giá ngoài thực tế, cần có camera thật, cảm biến thời tiết, dữ liệu lưu lượng và tốc độ thực tế, khoảng cách A/B được khảo sát cùng quy trình xác nhận trước khi gửi lệnh đến biển báo.")

for paragraph in doc.paragraphs:
    for run in paragraph.runs:
        if run.font.size is None:
            set_run_font(run, size=13)
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0

doc.save(FINAL)
print(FINAL)
