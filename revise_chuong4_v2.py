from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

from build_chuong4 import (
    OUT,
    add_block,
    add_figure_placeholder,
    add_heading,
    add_table,
    add_title,
    format_paragraph,
    set_run_font,
)

FINAL = "CHUONG_4_VSL_viet_lai_cap_nhat.docx"


def replace_paragraph(paragraph, text):
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)
    r = paragraph.add_run(text)
    set_run_font(r, size=13)


doc = Document(OUT)

# Update the chapter title and headings to the new single-chapter structure.
for p in doc.paragraphs:
    if p.text.startswith("CHƯƠNG 4."):
        replace_paragraph(p, "CHƯƠNG 4. KẾT QUẢ THỰC NGHIỆM, ĐÁNH GIÁ VÀ HƯỚNG PHÁT TRIỂN")
    elif p.text.startswith("4.1."):
        replace_paragraph(p, "4.1. Môi trường cài đặt và công nghệ sử dụng")
    elif p.text.startswith("4.3."):
        replace_paragraph(p, "4.3. Kết quả thực nghiệm và kịch bản thử nghiệm")

# Add the required English expansions and the JSON/history details to the existing sections.
for p in doc.paragraphs:
    if p.text.startswith("Chương này trình bày phần cài đặt"):
        replace_paragraph(p, "Chương này trình bày kết quả thực nghiệm, đánh giá hệ thống, thảo luận, hạn chế, kết luận và hướng phát triển của hệ thống giám sát giao thông thông minh có đề xuất tốc độ giới hạn linh hoạt VSL (Variable Speed Limit). Sau khi các chương trước đã giới thiệu lý do chọn đề tài, mục tiêu, cơ sở lý thuyết và nguyên lý xử lý, nội dung ở đây tập trung vào sản phẩm phần mềm đã xây dựng, luồng hoạt động thực tế, các kịch bản mô phỏng, kết quả đạt được và những vấn đề còn tồn tại. Cách trình bày được tổ chức theo đúng quá trình người vận hành sử dụng hệ thống, từ đăng nhập, chọn dữ liệu đầu vào, theo dõi phương tiện cho đến xem kết quả VSL.")
    elif p.text.startswith("Hệ thống được xây dựng ở mức mô phỏng"):
        replace_paragraph(p, "Hệ thống được xây dựng ở mức phần mềm mô phỏng trên máy tính cá nhân. Nguồn dữ liệu đầu vào gồm các video giao thông đã chuẩn bị sẵn hoặc camera được cấu hình trong chương trình. Khi một nguồn được chọn, chương trình đọc từng khung hình, đưa khung hình qua mô hình nhận diện, lọc đối tượng theo vùng giám sát ROI (Region of Interest), cập nhật số lượng phương tiện và tính các thông số phục vụ đề xuất tốc độ. Kết quả xử lý được hiển thị trực tiếp trên giao diện PyQt5 để người vận hành có thể quan sát cả hình ảnh lẫn các giá trị phân tích.")
    elif p.text.startswith("Python được sử dụng làm ngôn ngữ lập trình chính"):
        replace_paragraph(p, "Python được sử dụng làm ngôn ngữ lập trình chính vì phù hợp với việc kết hợp xử lý ảnh, mô hình học máy, giao diện desktop và lưu trữ dữ liệu trong cùng một chương trình. OpenCV đảm nhiệm việc mở video hoặc camera, đọc frame, chuyển đổi và vẽ thông tin lên khung hình như bounding box, nhãn phương tiện, ROI và hai vạch A/B. Mô hình YOLO (You Only Look Once) được dùng để phát hiện phương tiện trong từng frame. Kết quả của YOLO gồm loại phương tiện, tọa độ bounding box và độ tin cậy, sau đó được chuyển cho các bước đếm xe, phân loại mật độ và tính VSL.")
    elif p.text.startswith("PyQt5 đảm nhiệm phần giao diện"):
        replace_paragraph(p, "PyQt5 đảm nhiệm phần giao diện, bao gồm màn hình đăng nhập, khu vực hiển thị video, các ô thông tin về số xe, mật độ, thời tiết, sự cố, tỷ lệ xe nặng và tốc độ VSL. NumPy hỗ trợ các phép tính trên tọa độ, danh sách điểm của ROI, bộ nhớ đệm số xe và các thao tác xử lý dữ liệu số. SQLite được sử dụng để lưu tài khoản và lịch sử đăng nhập, đồng thời lưu một số dữ liệu cấu hình cần thiết trong phạm vi mô phỏng. JSON (JavaScript Object Notation) được dùng để lưu cấu hình camera, đa giác ROI và vị trí hai vạch đo. Dữ liệu kết quả có thể xuất thành CSV (Comma-Separated Values) để thuận tiện xử lý bằng bảng tính, đồng thời tạo HTML Report (HyperText Markup Language Report) để xem lại theo dạng báo cáo có cấu trúc. MQTT (Message Queuing Telemetry Transport) được đưa vào như một kênh mô phỏng việc gửi tốc độ VSL xuống biển báo điện tử.")
    elif p.text.startswith("Chức năng đăng nhập được đặt ở bước đầu"):
        replace_paragraph(p, "Chức năng đăng nhập được đặt ở bước đầu để kiểm soát quyền truy cập vào phần mềm. Người sử dụng cần nhập tài khoản và mật khẩu trước khi vào giao diện giám sát. Thông tin tài khoản và lịch sử đăng nhập được lưu trong cơ sở dữ liệu SQLite, nhờ đó chương trình có thể kiểm tra dữ liệu đăng nhập và ghi nhận phiên sử dụng thay vì để bất kỳ người dùng nào mở trực tiếp màn hình điều khiển. Trong bối cảnh VSL, việc kiểm soát người dùng có ý nghĩa vì tốc độ đề xuất và tốc độ nhập thủ công đều liên quan đến vận hành giao thông. Ở phạm vi đồ án, đăng nhập chưa phải là một hệ thống phân quyền nhiều cấp, nhưng đã tạo được bước kiểm soát cơ bản và giúp sản phẩm gần với một phần mềm vận hành thực tế hơn.")
    elif p.text.startswith("Sau bước chạy cơ bản, các tham số thời tiết"):
        replace_paragraph(p, p.text + " Sau mỗi phiên, số xe, mật độ, thời tiết, sự cố, tỷ lệ xe nặng, tốc độ VSL và thời gian phân tích được lưu để xuất CSV hoặc HTML Report.")

# Remove the old Chapter 4 tail, including the old limitations/development/location sections.
old_summary = doc.tables[2]
old_summary._tbl.getparent().remove(old_summary._tbl)
old_44 = next(p for p in doc.paragraphs if p.text.startswith("4.4. Hạn chế của hệ thống"))
body = doc._element.body
remove_started = False
for child in list(body):
    if child is old_44._p:
        remove_started = True
    if remove_started and child.tag != qn("w:sectPr"):
        body.remove(child)

add_heading(doc, "4.4. Đánh giá hệ thống và kết quả đạt được")
add_block(doc, """
Sau quá trình cài đặt và chạy thử, hệ thống đã mô phỏng được luồng xử lý cơ bản của một hệ thống giám sát giao thông thông minh có đề xuất VSL. Chương trình đã có sản phẩm chạy thử trên máy tính cá nhân, có giao diện để người vận hành đăng nhập, chọn nguồn dữ liệu, quan sát video và theo dõi các thông số phân tích. Hệ thống chưa phải là sản phẩm có thể triển khai ngay trên tuyến cao tốc, nhưng các thành phần chính đã được nối với nhau thành một quy trình có đầu vào, xử lý, kết quả và báo cáo.

Về chức năng đầu vào, hệ thống đọc được video giao thông và camera mô phỏng thông qua OpenCV. Các frame được đưa liên tục vào mô hình YOLO để nhận diện phương tiện. Kết quả nhận diện không chỉ được vẽ lên ảnh để minh họa mà còn được sử dụng tiếp cho các bước lọc ROI, đếm xe, phân loại phương tiện, tính tỷ lệ xe nặng và theo dõi xe đi qua hai vạch A/B. Điều này cho thấy phần thị giác máy tính đã đóng vai trò là bộ phận lấy dữ liệu từ video cho các bước phân tích phía sau.

Hệ thống xác định được ROI và chỉ đưa những phương tiện nằm trong vùng giám sát vào thống kê. Nhờ đó, số xe được đếm theo đúng khu vực cần quan sát hơn so với việc lấy toàn bộ khung hình. Số lượng xe trong ROI được làm ổn định bằng cửa sổ trượt rồi dùng để phân loại trạng thái giao thông. Các trạng thái như giao thông tốt, ổn định, mật độ cao và nguy cơ ùn tắc giúp thuật toán VSL có căn cứ điều chỉnh tốc độ thay vì chỉ dựa vào một giá trị đếm tức thời.

Chức năng đo tốc độ qua hai vạch A/B đã thể hiện được nguyên lý tính tốc độ từ video. Khi xe đi qua vạch A, hệ thống ghi thời điểm bắt đầu; khi cùng xe đi qua vạch B, hệ thống ghi thời điểm kết thúc và tính tốc độ theo khoảng cách giữa hai vạch. Trong điều kiện xe không bị che khuất và tracking ổn định, giao diện có thể hiển thị tốc độ ước lượng cạnh phương tiện. Kết quả này phù hợp với mục tiêu minh họa cách trích xuất tốc độ từ camera, nhưng chưa thể xem là phép đo hiện trường vì khoảng cách và phối cảnh chưa được khảo sát đầy đủ.

Phần tính VSL đã kết hợp mật độ, thời tiết, tỷ lệ xe nặng, sự cố và trạng thái giao thông. Khi số xe ít, thời tiết tốt và không có sự cố, tốc độ đề xuất giữ ở mức cao. Khi số xe tăng, trời mưa, sương mù, tỷ lệ truck/bus tăng hoặc sự cố xuất hiện, tốc độ được giảm thêm. Chế độ thủ công cho phép người vận hành nhập giá trị và giữ vai trò kiểm soát cuối cùng. Giao diện hiển thị video, ROI, bounding box, số xe, mật độ, thời tiết, sự cố, tốc độ VSL và nhật ký hệ thống, nên phù hợp cho mục đích demo và kiểm tra kết quả.

Về lưu trữ, hệ thống đã ghi nhận được dữ liệu sau phiên chạy bằng CSV và HTML Report. Các trường dữ liệu gồm số xe, mật độ, thời tiết, sự cố, tỷ lệ xe nặng, tốc độ VSL và thời gian phân tích. Việc lưu báo cáo giúp nhóm có thể đối chiếu kết quả giữa các kịch bản, thay vì chỉ dựa vào quan sát trực tiếp trên màn hình. SQLite cũng giúp duy trì thông tin tài khoản và lịch sử đăng nhập trong quá trình sử dụng.
""")

summary_rows = [
    ["Đọc video/camera", "Hệ thống đọc được video hoặc camera mô phỏng", "Tạo dữ liệu đầu vào cho xử lý"],
    ["Nhận diện phương tiện", "Phát hiện xe bằng YOLO", "Cung cấp dữ liệu cho đếm xe và VSL"],
    ["Xác định ROI", "Lọc xe theo vùng giám sát", "Giảm đếm nhầm ngoài khu vực phân tích"],
    ["Đếm xe", "Thống kê số xe trong ROI", "Làm cơ sở phân loại mật độ"],
    ["Phân loại mật độ", "Xác định trạng thái giao thông", "Hỗ trợ quyết định VSL"],
    ["Đo tốc độ", "Ước lượng tốc độ qua hai vạch A/B", "Bổ sung thông tin tốc độ phương tiện"],
    ["Tính VSL", "Đề xuất tốc độ theo mật độ, thời tiết, xe nặng và sự cố", "Thể hiện trọng tâm đề tài"],
    ["Giao diện", "Hiển thị kết quả trực quan", "Hỗ trợ người vận hành quan sát"],
    ["Điều khiển thủ công", "Cho phép nhập tốc độ thủ công", "Giữ vai trò kiểm soát của con người"],
    ["Lưu báo cáo", "Ghi nhận dữ liệu sau phiên chạy", "Phục vụ đánh giá kết quả thử nghiệm"],
]
add_table(doc, ["Nội dung đánh giá", "Kết quả đạt được", "Ý nghĩa"], summary_rows, [1.45, 3.15, 1.67], 9, "Bảng 4.3. Tổng hợp kết quả đạt được")

add_heading(doc, "4.5. Thảo luận kết quả")
add_block(doc, """
Kết quả thực nghiệm cho thấy hệ thống đã thể hiện đúng nguyên lý cơ bản của VSL: tốc độ giới hạn giảm khi mức độ rủi ro giao thông tăng. Với các kịch bản có ít xe, đường thông thoáng và thời tiết tốt, tốc độ đề xuất được giữ ở mức cao. Khi mật độ tăng hoặc xuất hiện các yếu tố bất lợi, tốc độ được điều chỉnh xuống. Mối quan hệ này phù hợp với mục tiêu ban đầu của đề tài là sử dụng dữ liệu từ video để hỗ trợ điều hành tốc độ thay đổi theo ngữ cảnh.

Trong các yếu tố đầu vào, mật độ phương tiện là nền tảng quan trọng nhất của thuật toán. Số xe trung bình trong ROI phản ánh trực tiếp mức độ chiếm dụng mặt đường và khoảng cách giữa các xe. Khi số xe tăng từ 3 lên 8, 13 và 18 xe, VSL lần lượt giảm từ 91 km/h xuống 76 km/h, 56 km/h và 40 km/h. Các giá trị này được xây dựng theo luật mô phỏng, nhưng xu hướng giảm là hợp lý vì dòng xe đông thường làm khoảng cách an toàn giảm, khả năng phanh gấp tăng và tốc độ chung khó duy trì ở mức cao.

Thời tiết và sự cố là những yếu tố làm thay đổi VSL rõ rệt ngoài ảnh hưởng của mật độ. Mưa nhỏ, mưa vừa và mưa to làm tốc độ giảm dần do mặt đường trơn, tầm nhìn kém và quãng đường phanh có thể tăng. Sương mù làm tốc độ giảm mạnh hơn trong một số kịch bản vì ảnh hưởng trực tiếp đến khả năng quan sát xe phía trước và chướng ngại vật. Khi có sự cố, thuật toán giảm tốc mạnh hơn để cảnh báo từ xa, đặc biệt ở trường hợp sự cố nghiêm trọng khi VSL được đưa về mức tối thiểu trong mô phỏng.

Tỷ lệ xe nặng giúp thuật toán gần với thực tế hơn vì không phải mọi phương tiện đều có ảnh hưởng giống nhau đến dòng xe. Xe tải và xe buýt thường có khối lượng lớn, quãng đường phanh dài hơn, chiếm nhiều không gian và có thể che khuất tầm nhìn của các xe phía sau. Khi tỷ lệ xe nặng tăng từ 15% lên 25% và 40%, tốc độ VSL giảm theo các mức đã thiết lập. Dù mức giảm hiện tại chưa được hiệu chỉnh bằng số liệu hiện trường, việc đưa yếu tố này vào mô hình là cần thiết để quyết định không chỉ dựa trên số lượng phương tiện.

Chế độ thủ công cũng có ý nghĩa trong bài toán an toàn giao thông. AI chỉ hỗ trợ lấy dữ liệu và đưa ra đề xuất, không nên tự quyết hoàn toàn trong mọi tình huống. Khi người vận hành nhập một tốc độ cụ thể, hệ thống ưu tiên tốc độ thủ công để có thể xử lý các tình huống đặc biệt như sự cố chưa được mô hình nhận diện, yêu cầu điều tiết tạm thời hoặc dữ liệu camera đang không ổn định. Trong triển khai thực tế, thao tác này cần gắn với xác nhận, phân quyền và nhật ký để bảo đảm trách nhiệm vận hành rõ ràng.

Nhìn chung, kết quả hiện tại có ý nghĩa ở mức mô phỏng và kiểm tra logic. Các giá trị 91 km/h, 56 km/h, 46 km/h hay 40 km/h cho thấy thuật toán thay đổi theo kịch bản, nhưng chưa thể xem là số liệu hiện trường hoặc tốc độ pháp lý áp dụng cho một tuyến cụ thể. Để đưa hệ thống ra thực tế, cần có camera thật, cảm biến thời tiết thật, dữ liệu lưu lượng và tốc độ thực tế, khoảng cách A/B được khảo sát, cùng quy trình xác nhận trước khi gửi lệnh đến biển báo.
""")

add_heading(doc, "4.6. Hạn chế của hệ thống")
add_block(doc, """
Hạn chế đầu tiên nằm ở chất lượng video đầu vào. Kết quả nhận diện phụ thuộc vào độ phân giải, ánh sáng, góc quay, độ rung của camera và kích thước phương tiện trong khung hình. Khi video bị tối, mờ, rung hoặc có mưa và sương mù, đường biên của xe không rõ, độ tin cậy của YOLO có thể giảm. Xe bị che khuất bởi xe khác, dải phân cách hoặc vật thể phía trước cũng làm detector bỏ sót xe hoặc tạo bounding box không ổn định.

Mô hình YOLO hiện tại chưa được huấn luyện riêng cho dữ liệu cao tốc Việt Nam. Vì vậy, hệ thống có thể nhầm giữa xe con, xe tải, xe buýt và xe khách, nhất là khi phương tiện ở xa hoặc chỉ nhìn thấy một phần. Sai khác trong phân loại làm ảnh hưởng đến tỷ lệ xe nặng và có thể khiến mức giảm VSL chưa phù hợp. Các lớp car, truck, bus, motorbike và bicycle mới phản ánh nhóm phương tiện cơ bản, chưa bao quát hết các biến thể như xe container, xe tải nhỏ hoặc xe khách nhiều kích thước.

ROI vẫn cần được chỉnh thủ công theo từng camera hoặc video. Nếu ROI quá rộng, xe ở làn ngược chiều, lề đường hoặc vùng không liên quan có thể bị đếm nhầm. Nếu ROI quá hẹp, các xe ở mép làn hoặc khu vực bị biến dạng phối cảnh có thể bị bỏ sót. Khi camera thay đổi góc nhìn, cấu hình ROI cũ cũng có thể không còn chính xác. Đây là lý do việc lưu ROI theo từng camera và nghiên cứu gợi ý ROI tự động là hướng cần thiết.

Đo tốc độ bằng hai vạch A/B chỉ cho tốc độ ước lượng. Công thức v = s / t × 3,6 yêu cầu biết khoảng cách thực tế giữa hai vạch và xác định đúng thời điểm cùng một xe đi qua A và B. Trong đồ án, khoảng cách này chưa được khảo sát ngoài hiện trường và ảnh chưa được hiệu chỉnh phối cảnh, nên tốc độ pixel trên khung hình chưa phản ánh chính xác khoảng cách ngoài đường. Nếu tracking mất dấu do xe bị che khuất hoặc hai xe đi quá gần nhau, thời điểm A và B có thể bị gán sai, làm kết quả tốc độ sai.

Thời tiết hiện chủ yếu do người vận hành chọn trên giao diện. Hệ thống chưa tự động nhận diện mưa, sương mù hoặc tầm nhìn từ video, cũng chưa kết nối cảm biến lượng mưa, độ ẩm, tầm nhìn và tình trạng mặt đường. Cách chọn thủ công phù hợp để kiểm tra có kiểm soát, nhưng khi vận hành liên tục sẽ phụ thuộc vào sự quan sát và thao tác của con người.

Thuật toán VSL được xây dựng theo các luật và ngưỡng giảm tốc mang tính mô phỏng. Các mức giảm do mật độ, thời tiết, xe nặng và sự cố chưa được hiệu chỉnh bằng dữ liệu thực tế của tuyến. Hệ thống hiện chủ yếu phản ứng với trạng thái đang quan sát, chưa có mô hình dự báo ùn tắc hoặc dự báo tốc độ trong vài phút tiếp theo. Do đó, kết quả phù hợp để đánh giá logic nhưng chưa đủ cơ sở dùng làm lệnh điều hành thật.

Về triển khai, hệ thống chưa kết nối camera thật, cảm biến thật và biển báo LED/VMS thật. MQTT mới dừng ở mức mô phỏng gửi giá trị VSL, chưa kiểm tra đầy đủ các tình huống mất mạng, xác thực thiết bị, gửi lặp, phản hồi trạng thái hoặc đưa biển báo về chế độ an toàn khi có lỗi. Giao diện hiện cũng chưa có đầy đủ biểu đồ lịch sử, trạng thái từng camera, cảnh báo sự cố và lịch sử thay đổi VSL để phục vụ một trung tâm điều hành nhiều tuyến.
""")

add_heading(doc, "4.7. Hướng phát triển và giải pháp cải thiện")
add_block(doc, """
Để cải thiện dữ liệu đầu vào, cần thu thập video ở nhiều điều kiện gồm ban ngày, ban đêm, trời quang, mưa nhỏ, mưa to, sương mù, mật độ thấp và mật độ cao. Camera nên được đặt ở vị trí có góc nhìn rõ mặt đường, bao quát các làn cần giám sát và ít bị che khuất. Ngoài chất lượng hình ảnh, cần ghi nhận FPS, thời điểm quay, vị trí camera và điều kiện thời tiết để dữ liệu sau này có thể dùng cho hiệu chỉnh thuật toán.

Mô hình YOLO nên được huấn luyện hoặc tinh chỉnh bằng dữ liệu giao thông Việt Nam. Bộ dữ liệu cần bổ sung xe khách, xe container, xe tải nhỏ, xe tải nặng và các phương tiện có hình dạng khác nhau. Dữ liệu cần có nhiều góc quay, khoảng cách, điều kiện ánh sáng và mức độ che khuất để mô hình phân biệt car, truck và bus tốt hơn. Khi độ chính xác loại xe được cải thiện, tỷ lệ xe nặng và mức giảm VSL sẽ có cơ sở đáng tin cậy hơn.

ROI nên được lưu riêng cho từng camera trong tệp JSON cùng với vị trí vạch A/B, hướng di chuyển và khoảng cách thực tế. Khi camera được chọn, phần mềm có thể tự nạp cấu hình đã lưu, giảm thao tác chỉnh thủ công. Về lâu dài, có thể nghiên cứu phát hiện làn đường, mặt đường hoặc vùng phương tiện xuất hiện ổn định để gợi ý ROI tự động cho người vận hành xác nhận.

Đối với đo tốc độ, cần khảo sát khoảng cách thực tế giữa hai vạch A và B tại từng vị trí. Camera nên được hiệu chỉnh phối cảnh để chuyển tọa độ ảnh sang mặt phẳng đường gần đúng hơn. Phần tracking có thể nâng cấp bằng DeepSORT, ByteTrack hoặc OC-SORT để giữ ID xe ổn định và giảm lỗi khi phương tiện đi qua các frame liên tiếp. Khi có mất dấu, hệ thống nên đánh dấu kết quả không chắc chắn thay vì chỉ hiển thị một tốc độ như thể đó là giá trị chính xác.

Thời tiết có thể được tự động hóa bằng mô hình AI nhận diện mưa, sương mù và tầm nhìn từ ảnh camera. Một hướng khác là kết nối cảm biến hoặc API thời tiết để lấy lượng mưa, độ ẩm, tầm nhìn và tình trạng mặt đường. Khi có nhiều nguồn dữ liệu, hệ thống có thể so sánh và xác định mức thời tiết ổn định hơn so với chỉ cho người vận hành chọn thủ công.

Thuật toán VSL cần được hiệu chỉnh bằng dữ liệu thật của tuyến đường. Có thể phân tích quan hệ giữa mật độ, tốc độ trung bình, tỷ lệ xe nặng, thời tiết, sự cố và các tình huống nguy hiểm để điều chỉnh ngưỡng. Khi dữ liệu đủ lớn, hệ thống có thể kết hợp luật cố định với mô hình học máy, đồng thời bổ sung mô hình dự báo mật độ hoặc tốc độ trong vài phút tiếp theo. Dù dùng phương pháp nào, tốc độ cuối cùng vẫn phải bị giới hạn bởi các mức vận hành được phê duyệt.

Giao diện có thể phát triển thành dashboard web để theo dõi nhiều camera, bản đồ tuyến, lịch sử VSL, cảnh báo sự cố và trạng thái kết nối. Dashboard giúp người vận hành nhìn được quan hệ giữa các camera CAM_01, CAM_02 và CAM_03, đồng thời xem lại sự thay đổi tốc độ theo thời gian. Khi mở rộng nhiều tuyến, giao diện web sẽ thuận tiện hơn việc chỉ chạy một màn hình desktop.

Hệ thống có thể kết nối biển báo LED/VMS thật nhưng cần cơ chế xác nhận của người vận hành trước khi gửi lệnh. Quy trình nên hiển thị tốc độ đề xuất, lý do, dữ liệu đầu vào và mức độ tin cậy; sau đó người vận hành xác nhận mới gửi lệnh qua MQTT hoặc giao thức tương ứng. Cách thiết kế này giữ vai trò hỗ trợ ra quyết định của AI và không thay thế hoàn toàn con người trong bài toán an toàn giao thông.
""")

add_heading(doc, "4.8. Đề xuất vị trí đặt biển báo VSL")
add_block(doc, """
Nguyên tắc đặt biển VSL là phải đặt đủ xa trước khu vực nguy hiểm để người lái có thời gian nhận biết và giảm tốc. Biển không nên đặt quá gần vị trí ùn tắc hoặc sự cố vì khi đó người lái có thể phải phanh gấp. Ngược lại, nếu đặt quá xa mà không có biển nhắc lại, thông tin về tốc độ có thể bị quên trước khi xe đi đến vùng nguy hiểm. Khoảng cách đặt biển cần được xem xét cùng tốc độ khai thác, tầm nhìn, độ cong, độ dốc, lưu lượng xe và khoảng cách phanh.

Các vị trí nên ưu tiên gồm đầu tuyến cao tốc, trước nút giao, trước lối ra hoặc lối vào cao tốc, trước trạm thu phí, trước đoạn cầu, trước đoạn cong, trước khu vực tầm nhìn hạn chế, khu vực thường xuyên ùn tắc, khu vực hay có mưa lớn hoặc sương mù và khu vực có camera phát hiện sự cố. Các vị trí này đều có khả năng làm dòng xe thay đổi tốc độ hoặc làm giảm thời gian phản ứng của người lái.

Theo mô hình camera mô phỏng, CAM_01 có thể đặt tại khoảng KM10, CAM_02 tại KM15 và CAM_03 tại KM20. Biển VSL được đề xuất đặt trước vùng giám sát từ 500 đến 1000 m tùy điều kiện mặt đường và tầm nhìn. Camera phân tích tình trạng giao thông phía trước, còn biển báo đặt ở phía trước vùng nguy hiểm để người lái nhận thông tin sớm và điều chỉnh tốc độ trước khi đi vào đoạn có mật độ cao hoặc sự cố.
""")

location_rows = [
    ["Đầu tuyến cao tốc", "Xe bắt đầu vào tuyến và cần biết tốc độ khai thác hiện tại", "Thông báo tốc độ phù hợp ngay từ đầu"],
    ["Trước nút giao", "Có nhập làn, tách làn và chuyển làn", "Giảm xung đột giữa các dòng xe"],
    ["Trước lối ra/vào cao tốc", "Lưu lượng thay đổi, xe thường chuyển làn", "Ổn định tốc độ trước vùng giao cắt"],
    ["Trước trạm thu phí", "Xe thường giảm tốc hoặc dừng theo làn", "Hạn chế phanh gấp và va chạm phía sau"],
    ["Trước đoạn cầu", "Không gian xử lý và vị trí dừng xe hạn chế", "Tăng an toàn khi mật độ hoặc thời tiết xấu"],
    ["Trước đoạn cong", "Tầm nhìn và khả năng quan sát phía trước giảm", "Giảm nguy cơ mất lái và vào cua quá nhanh"],
    ["Khu vực tầm nhìn hạn chế", "Người lái khó nhận biết xe dừng hoặc chướng ngại vật", "Tạo cảnh báo sớm bằng tốc độ thấp hơn"],
    ["Khu vực thường xuyên ùn tắc", "Dòng xe giảm tốc bất thường theo thời điểm", "Điều hòa tốc độ từ xa, giảm dồn xe"],
    ["Khu vực hay mưa lớn/sương mù", "Mặt đường trơn hoặc tầm nhìn giảm", "Điều chỉnh tốc độ theo điều kiện quan sát"],
    ["Trước vùng CAM_01/CAM_02/CAM_03 tại KM10/KM15/KM20", "Camera phát hiện mật độ, xe nặng hoặc sự cố phía trước", "Đặt biển trước vùng giám sát khoảng 500–1000 m để xe kịp giảm tốc"],
]
add_table(doc, ["Nhóm vị trí", "Lý do đặt biển VSL", "Mục tiêu"], location_rows, [1.60, 2.50, 2.10], 9, "Bảng 4.4. Đề xuất nhóm vị trí đặt biển báo VSL")
add_block(doc, """
Khi có ùn tắc hoặc sự cố nghiêm trọng, có thể bố trí nhiều biển liên tiếp để xe giảm tốc theo chuỗi 100 → 80 → 60 → 40 km/h. Biển đầu tiên thông báo giảm tốc ở khoảng cách xa, các biển tiếp theo nhắc lại và đưa tốc độ xuống gần vùng nguy hiểm. Cách giảm tốc theo chuỗi giúp dòng xe giảm dần, hạn chế phanh gấp và giảm nguy cơ va chạm phía sau. Khi sự cố được xử lý, tốc độ cũng nên được khôi phục theo từng mức và có xác nhận của người vận hành.
""")
add_figure_placeholder(doc, 9, "Đề xuất vị trí đặt biển báo VSL theo cụm camera trên tuyến", "sơ đồ tuyến hoặc hình minh họa camera CAM_01, CAM_02, CAM_03 và biển báo VSL đặt trước vùng giám sát.")
add_figure_placeholder(doc, 10, "Minh họa giảm tốc theo chuỗi biển báo VSL", "sơ đồ các biển báo liên tiếp hiển thị 100 km/h, 80 km/h, 60 km/h, 40 km/h trước khu vực ùn tắc hoặc sự cố.")

add_heading(doc, "4.9. Kết luận chương")
add_block(doc, """
Chương 4 đã trình bày toàn bộ kết quả của phần cài đặt, thực nghiệm, đánh giá, thảo luận, hạn chế và hướng phát triển của hệ thống. Việc gộp nội dung theo một chương duy nhất giúp thể hiện liền mạch từ sản phẩm phần mềm, các kịch bản kiểm thử, kết quả đạt được cho đến những vấn đề cần tiếp tục nghiên cứu. Hệ thống được xây dựng bằng Python, OpenCV, YOLO, PyQt5, SQLite, JSON, CSV/HTML Report và MQTT, trong đó mỗi công nghệ đảm nhiệm một vai trò cụ thể trong luồng xử lý.

Kết quả triển khai cho thấy hệ thống đọc được video hoặc camera, nhận diện phương tiện, hiển thị bounding box và nhãn xe, lọc phương tiện theo ROI, đếm xe, phân loại mật độ, đo tốc độ ước lượng qua hai vạch A/B, tính VSL theo mật độ, thời tiết, tỷ lệ xe nặng và sự cố, hiển thị giao diện, hỗ trợ điều khiển thủ công và lưu báo cáo. Đây là các chức năng chính cần có trong mô hình giám sát giao thông và đề xuất tốc độ ở mức mô phỏng.

Qua 20 kịch bản thử nghiệm, VSL giảm khi mật độ tăng, giảm khi thời tiết xấu, giảm khi tỷ lệ xe nặng cao và giảm mạnh khi có sự cố. Chế độ thủ công giúp người vận hành kiểm soát hệ thống, phù hợp với quan điểm AI chỉ hỗ trợ ra quyết định chứ không thay thế hoàn toàn con người. Kết quả này thể hiện đúng nguyên lý VSL, nhưng các con số vẫn cần được hiểu là kết quả của bộ ngưỡng mô phỏng.

Hệ thống còn phụ thuộc vào chất lượng video, mô hình YOLO, cấu hình ROI, khả năng tracking, độ chính xác của khoảng cách A/B và cách chọn thời tiết. Tốc độ đo được mới là ước lượng, thời tiết chưa tự động nhận diện, VSL chưa hiệu chỉnh bằng dữ liệu thực tế và MQTT chưa kết nối với thiết bị ngoài hiện trường. Vì vậy, hệ thống chưa thể xem là giải pháp vận hành hoàn chỉnh.

Hướng phát triển gồm huấn luyện YOLO bằng dữ liệu giao thông Việt Nam, nâng cấp tracking, hiệu chỉnh phối cảnh, tự động nhận diện thời tiết, hiệu chỉnh thuật toán VSL bằng dữ liệu thật, xây dựng dashboard web và kết nối biển báo LED/VMS có xác nhận của người vận hành. Với phạm vi đã thực hiện, hệ thống đạt mục tiêu ở mức mô phỏng, phù hợp với hướng Trí tuệ nhân tạo và Giao thông thông minh, đồng thời có nền tảng để tiếp tục mở rộng.
""")

add_title(doc, "KẾT LUẬN CHUNG")
add_block(doc, """
Đề tài “Xây dựng hệ thống giám sát giao thông thông minh và đề xuất tốc độ giới hạn linh hoạt VSL dựa trên thị giác máy tính” được thực hiện nhằm mô phỏng một hệ thống có khả năng thu nhận dữ liệu giao thông từ video hoặc camera, phân tích tình trạng phương tiện và hỗ trợ đề xuất tốc độ giới hạn theo từng tình huống. Ý tưởng chính của đề tài là sử dụng dữ liệu quan sát được để thay đổi tốc độ theo mật độ, thời tiết, tỷ lệ xe nặng, sự cố và trạng thái giao thông thay vì chỉ áp dụng một mức tốc độ cố định.

Trong quá trình thực hiện, nhóm đã xây dựng được phần mềm có khả năng đăng nhập, chọn video hoặc camera, nhận diện phương tiện bằng YOLO, xác định ROI, đếm xe, phân loại mật độ, đo tốc độ ước lượng qua hai vạch A/B và tính VSL. Kết quả được hiển thị trên giao diện PyQt5, lưu tài khoản và lịch sử đăng nhập bằng SQLite, lưu cấu hình camera và ROI bằng JSON, xuất dữ liệu bằng CSV/HTML Report và mô phỏng hướng truyền tốc độ đến biển báo bằng MQTT.

Về Trí tuệ nhân tạo (Artificial Intelligence - AI), đề tài đã sử dụng YOLO để nhận diện phương tiện từ video. Kết quả nhận diện gồm loại xe, bounding box và độ tin cậy được dùng tiếp cho đếm xe, tính tỷ lệ xe nặng, phân loại mật độ và tạo đầu vào cho thuật toán VSL. Như vậy, AI trong đề tài không chỉ có vai trò minh họa trên khung hình mà còn tham gia vào quá trình tạo dữ liệu phục vụ quyết định tốc độ.

Về Giao thông thông minh, đề tài đã mô phỏng được quy trình thu thập dữ liệu, phân tích trạng thái giao thông và hỗ trợ điều hành tốc độ. Kết quả thử nghiệm cho thấy khi mật độ tăng, thời tiết xấu, xe nặng nhiều hoặc có sự cố, VSL giảm theo hướng an toàn hơn. Khi đường thông thoáng và điều kiện quan sát tốt, VSL có thể duy trì ở mức cao hơn. Đây là kết quả phù hợp với mục tiêu nghiên cứu của đề tài.

Tuy nhiên, hệ thống vẫn còn hạn chế về chất lượng video, dữ liệu huấn luyện YOLO, cấu hình ROI, tracking, hiệu chỉnh tốc độ A/B, nhận diện thời tiết và các ngưỡng VSL. Hệ thống chưa kết nối camera thật, cảm biến thật và biển báo LED/VMS thật, nên chưa thể thay thế một hệ thống điều hành giao thông ngoài thực tế. Các hạn chế này cần được xem là cơ sở để đánh giá đúng phạm vi của đồ án, không phải là những kết quả đã được kiểm chứng trên tuyến đường.

Trong tương lai, đề tài có thể tiếp tục phát triển bằng dữ liệu giao thông Việt Nam, mô hình YOLO được huấn luyện riêng, tracking ổn định hơn, hiệu chỉnh phối cảnh, nhận diện thời tiết tự động, dự báo mật độ trong vài phút tiếp theo, dashboard web cho nhiều camera và kết nối biển báo thật theo cơ chế xác nhận. Nhìn chung, đồ án đã hoàn thành mục tiêu ở mức mô phỏng và thể hiện được sự kết hợp giữa Trí tuệ nhân tạo, xử lý ảnh, lập trình giao diện và Giao thông thông minh. Đây là nền tảng phù hợp để tiếp tục nghiên cứu một hệ thống VSL có độ tin cậy cao hơn trong tương lai.
""")

# Reapply spacing/font normalization after paragraph replacement and appended content.
for paragraph in doc.paragraphs:
    for run in paragraph.runs:
        if run.font.size is None:
            set_run_font(run, size=13)
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0

doc.save(FINAL)
print(FINAL)
