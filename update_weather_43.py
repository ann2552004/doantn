from copy import deepcopy

from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn

from build_chuong4 import add_block, add_table, format_paragraph, set_run_font


SOURCE = "CHUONG_4_VSL_viet_lai_cap_nhat.docx"
FINAL = "CHUONG_4_VSL_cap_nhat_thoi_tiet.docx"


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(13)


def replace_paragraph(paragraph, text):
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)
    r = paragraph.add_run(text)
    set_run_font(r, size=13)


def weather_fragment():
    tmp = Document()
    configure_document(tmp)
    rows = [
        [1, "Trời quang", "Chọn trạng thái thời tiết bình thường, mặt đường khô, tầm nhìn tốt", "VSL giữ ở mức cao nếu mật độ xe thấp", "Làm kịch bản nền để so sánh với các điều kiện thời tiết xấu"],
        [2, "Trời mưa", "Chọn trạng thái mưa trên giao diện", "VSL giảm so với trời quang", "Kiểm tra khả năng hệ thống giảm tốc khi mặt đường trơn và quãng đường phanh tăng"],
        [3, "Sương mù", "Chọn trạng thái sương mù trên giao diện", "VSL giảm mạnh hơn so với trời quang và có thể thấp hơn mưa", "Kiểm tra phản ứng của hệ thống khi tầm nhìn bị hạn chế"],
        [4, "Thời tiết xấu kết hợp mật độ cao", "Chạy video có nhiều xe trong ROI, đồng thời chọn mưa hoặc sương mù", "VSL giảm rõ rệt hơn so với trường hợp chỉ có mật độ cao hoặc chỉ có thời tiết xấu", "Kiểm tra cách thuật toán kết hợp nhiều yếu tố rủi ro cùng lúc"],
    ]
    add_table(tmp, ["STT", "Kịch bản thời tiết", "Điều kiện kiểm thử", "Kết quả mong đợi", "Ý nghĩa đánh giá"], rows, [0.35, 1.05, 1.75, 1.65, 1.40], 8, "Bảng 4.1. Các kịch bản kiểm thử thời tiết trong hệ thống VSL")
    add_block(tmp, """
Các kịch bản thời tiết được chọn theo hướng đại diện thay vì tách quá nhiều mức nhỏ. Trời quang được dùng làm điều kiện nền vì đây là trạng thái mặt đường và tầm nhìn thuận lợi nhất. Khi chạy ở điều kiện này, nếu mật độ phương tiện thấp và không có sự cố, tốc độ VSL có thể được giữ ở mức cao hơn.

Với kịch bản trời mưa, hệ thống cần giảm tốc độ đề xuất so với trời quang. Nguyên nhân là khi mưa, mặt đường trơn hơn, độ bám giữa lốp xe và mặt đường giảm, đồng thời quãng đường phanh của phương tiện có thể tăng. Vì vậy, việc giảm VSL trong điều kiện mưa giúp cảnh báo người lái di chuyển thận trọng hơn.

Với kịch bản sương mù, hệ thống cần giảm tốc mạnh hơn vì tầm nhìn của người lái bị hạn chế. Khi người lái khó quan sát xe phía trước, vật cản hoặc sự thay đổi tốc độ của dòng xe, việc duy trì tốc độ cao sẽ làm tăng nguy cơ va chạm. Do đó, sương mù được xem là điều kiện thời tiết có mức rủi ro cao hơn mưa trong hệ thống mô phỏng.

Kịch bản thời tiết xấu kết hợp mật độ cao được dùng để kiểm tra phản ứng của thuật toán khi nhiều yếu tố bất lợi xuất hiện cùng lúc. Trong trường hợp vừa có nhiều phương tiện trong ROI, vừa có mưa hoặc sương mù, hệ thống cần giảm VSL rõ rệt hơn so với trường hợp chỉ có một yếu tố riêng lẻ. Kịch bản này phù hợp với thực tế vì các tình huống nguy hiểm trên đường thường không chỉ đến từ một nguyên nhân duy nhất.

Kết quả kiểm thử thời tiết trong đồ án được hiểu là kết quả mô phỏng để đánh giá logic của thuật toán VSL. Các mức tốc độ đề xuất chưa được xem là tốc độ pháp lý áp dụng cho tuyến đường thực tế. Khi triển khai ngoài hiện trường, các mức giảm tốc do thời tiết cần được hiệu chỉnh bằng dữ liệu thực tế, cảm biến thời tiết hoặc dữ liệu tầm nhìn tại tuyến.
""")
    return tmp


doc = Document(SOURCE)
configure_document(doc)
fragment = weather_fragment()

# Replace the old general/20-case scenario block with the representative weather block.
start = next(p for p in doc.paragraphs if p.text.startswith("Bảng 4.1. Các kịch bản thử nghiệm chính"))
end = next(p for p in doc.paragraphs if p.text.startswith("Về kết quả đạt được, hệ thống đã đọc được video giao thông"))
body = doc._element.body
removing = False
for child in list(body):
    if child == start._p:
        removing = True
    if removing and child != end._p and child.tag != qn("w:sectPr"):
        body.remove(child)
    if child == end._p:
        break

for child in list(fragment._element.body):
    if child.tag != qn("w:sectPr"):
        end._p.addprevious(deepcopy(child))

# Keep one caption for the retained summary table and renumber the following location table.
captions = [p for p in doc.paragraphs if p.text.startswith("Bảng 4.3. Tổng hợp kết quả đạt được")]
if captions:
    replace_paragraph(captions[0], "Bảng 4.2. Tổng hợp kết quả đạt được")
    for extra in captions[1:]:
        extra._element.getparent().remove(extra._element)
for p in doc.paragraphs:
    if p.text.startswith("Bảng 4.4. Đề xuất nhóm vị trí đặt biển báo VSL"):
        replace_paragraph(p, "Bảng 4.3. Đề xuất nhóm vị trí đặt biển báo VSL")

# Keep 4.5 and the chapter conclusion consistent with representative weather groups.
for p in doc.paragraphs:
    if p.text.startswith("Thời tiết và sự cố là những yếu tố làm thay đổi VSL rõ rệt"):
        replace_paragraph(p, "Thời tiết và sự cố là những yếu tố làm thay đổi VSL rõ rệt ngoài ảnh hưởng của mật độ. Trời mưa làm mặt đường trơn hơn và có thể làm tăng quãng đường phanh, trong khi sương mù làm giảm tầm nhìn nên hệ thống có thể giảm tốc mạnh hơn. Khi thời tiết xấu xuất hiện đồng thời với mật độ cao hoặc sự cố, VSL cần được điều chỉnh thấp hơn để cảnh báo dòng xe phía sau. Các mức điều chỉnh hiện tại vẫn mang tính mô phỏng và cần được kiểm chứng bằng dữ liệu thực tế.")
    elif p.text.startswith("Để cải thiện dữ liệu đầu vào, cần thu thập video ở nhiều điều kiện"):
        replace_paragraph(p, "Để cải thiện dữ liệu đầu vào, cần thu thập video ở nhiều điều kiện gồm ban ngày, ban đêm, trời quang, trời mưa, sương mù, mật độ thấp và mật độ cao. Camera nên được đặt ở vị trí có góc nhìn rõ mặt đường, bao quát các làn cần giám sát và ít bị che khuất. Ngoài chất lượng hình ảnh, cần ghi nhận FPS, thời điểm quay, vị trí camera và điều kiện thời tiết để dữ liệu sau này có thể dùng cho hiệu chỉnh thuật toán.")
    elif p.text.startswith("Qua 20 kịch bản thử nghiệm"):
        replace_paragraph(p, "Qua các kịch bản kiểm thử, VSL giảm khi mật độ tăng, giảm khi trời mưa hoặc sương mù, giảm khi tỷ lệ xe nặng cao và giảm mạnh khi có sự cố. Chế độ thủ công giúp người vận hành kiểm soát hệ thống, phù hợp với quan điểm AI chỉ hỗ trợ ra quyết định chứ không thay thế hoàn toàn con người. Kết quả này thể hiện đúng xu hướng của VSL, nhưng vẫn cần được hiểu là kết quả mô phỏng theo các tham số đã cài đặt.")

for paragraph in doc.paragraphs:
    for run in paragraph.runs:
        if run.font.size is None:
            set_run_font(run, size=13)
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0

doc.save(FINAL)
print(FINAL)
